import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_document():
    doc = Document()
    
    # Page setup - Margins (1 inch all sides)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Color Palette - Professional B.Tech Academic Theme
    PRIMARY_COLOR = RGBColor(16, 185, 129)    # Emerald Green #10B981
    SECONDARY_COLOR = RGBColor(245, 158, 11)  # Trophy Gold #F59E0B
    DARK_BG = RGBColor(15, 23, 42)            # Slate Dark #0F172A
    TEXT_COLOR = RGBColor(30, 41, 59)          # Slate Text #1E293B
    MUTED_COLOR = RGBColor(100, 116, 139)      # Slate Muted #64748B

    def set_cell_background(cell, fill_hex):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
            node = OxmlElement(margin)
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = TEXT_COLOR

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(17)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(13.5)
        run.font.bold = True
        run.font.color.rgb = SECONDARY_COLOR
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = DARK_BG
        return p

    def add_bullet(bold_prefix="", text=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.bold = True
        p.add_run(text)
        return p

    def add_code(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.name = 'Consolas'
        r.font.size = Pt(9)
        r.font.color.rgb = DARK_BG
        return p

    def style_table(table, col_widths, headers, data):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        for i, header_text in enumerate(headers):
            hdr_cells[i].text = header_text
            set_cell_background(hdr_cells[i], "0F172A")
            set_cell_margins(hdr_cells[i], top=120, bottom=120, left=140, right=140)
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
        
        for row_idx, row_data in enumerate(data):
            row_cells = table.add_row().cells
            bg_color = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
            for i, val in enumerate(row_data):
                row_cells[i].text = str(val)
                set_cell_background(row_cells[i], bg_color)
                set_cell_margins(row_cells[i], top=100, bottom=100, left=140, right=140)
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(9.5)
                        run.font.color.rgb = TEXT_COLOR
                        
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    # -------------------------------------------------------------
    # TITLE BLOCK
    # -------------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(2)
    t_run = title_p.add_run("ODI World Cup 2027 AI Analytics & Raj vs AI Prediction Platform")
    t_run.font.size = Pt(24)
    t_run.font.bold = True
    t_run.font.color.rgb = PRIMARY_COLOR

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(8)
    s_run = sub_p.add_run("AI/ML-Based ODI World Cup Historical Analysis and Winner Prediction System")
    s_run.font.size = Pt(14)
    s_run.font.bold = True
    s_run.font.color.rgb = SECONDARY_COLOR

    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_after = Pt(20)
    m_run = meta_p.add_run("Author: Raj Malviya  |  Technology: React 18 + Vite 5 + Tailwind CSS 3 + Python 3.10+ + FastAPI + Scikit-learn")
    m_run.font.size = Pt(10)
    m_run.font.italic = True
    m_run.font.color.rgb = MUTED_COLOR

    # -------------------------------------------------------------
    # 1. EXECUTIVE SUMMARY
    # -------------------------------------------------------------
    add_h1("1. Executive Summary")
    p1 = doc.add_paragraph()
    p1.add_run("The ")
    p1.add_run("ODI World Cup 2027 AI Analytics & Raj vs AI Prediction Platform").bold = True
    p1.add_run(" is an enterprise-grade full-stack data analytics and machine learning application developed to evaluate 40 years of ICC Men's ODI World Cup historical records (1983–2023), team strength indices, and regional venue adaptability metrics. The project introduces a quantitative alternative to traditional, subjective sports commentary by implementing a supervised machine learning pipeline evaluated via 5-Fold Stratified Cross-Validation across 4 classification algorithms.")

    p2 = doc.add_paragraph()
    p2.add_run("The primary objective of the system is to calculate Softmax-calibrated win probability distributions for the upcoming 2027 ODI World Cup (hosted jointly across South Africa, Zimbabwe, and Namibia) while delivering an interactive, gamified prediction comparison engine called ")
    p2.add_run("Raj vs AI").bold = True
    p2.add_run(". Through this feature, users can manually select their predicted winners for every World Cup edition from 1983 to 2027. The system dynamically evaluates both the user's picks and the machine learning model's predictions against actual historical champions, maintaining persistent client-side state in LocalStorage, updating live scoreboards, and tracking leader standings.")

    p3 = doc.add_paragraph()
    p3.add_run("Built using a modern microservice-inspired architecture, the application combines a high-performance ")
    p3.add_run("Python FastAPI").bold = True
    p3.add_run(" REST backend with a responsive ")
    p3.add_run("React 18 + Vite 5 + Tailwind CSS").bold = True
    p3.add_run(" single-page application. The selected machine learning model, ")
    p3.add_run("Logistic Regression").bold = True
    p3.add_run(" (Validation Accuracy: 74.38%, F1-Score: 0.5070, Historical Backtest Accuracy: 72.7%), predicts Australia as the leading contender for 2027 with a 31.3% win probability, closely followed by India at 23.9%.")

    # -------------------------------------------------------------
    # 2. VISION
    # -------------------------------------------------------------
    add_h1("2. Vision")
    p = doc.add_paragraph()
    p.add_run("The vision of the platform is to democratize sports analytics by providing fans, analysts, and researchers with transparent, quantitative tools that decouple tournament predictions from media speculation and emotional bias. By transforming multi-decade ICC statistics into 10 normalized feature vectors, the system establishes a standardized foundation for evaluating team performance.")

    p2 = doc.add_paragraph()
    p2.add_run("Looking forward, the platform aims to evolve from a tournament-level winner prediction model into a real-time sports intelligence ecosystem. Future iterations will incorporate live ball-by-ball match webhooks, player-level Elo ratings, atmospheric venue condition APIs, and Monte Carlo bracket simulations, creating an indispensable resource for academic sports data science.")

    # -------------------------------------------------------------
    # 3. BUSINESS PROBLEM
    # -------------------------------------------------------------
    add_h1("3. Business & Product Problem")
    p = doc.add_paragraph()
    p.add_run("Predicting the champion of a major multi-week international tournament like the ODI World Cup presents significant analytical challenges. Traditional commentary often fails due to the following specific product and domain problems:")

    add_bullet("1. Unstructured Historical Data: ", "Over 40 years of World Cup history comprising 556 matches remain locked in disparate scorecards, making multi-tournament trend comparison cumbersome.")
    add_bullet("2. Subjective Fan & Media Bias: ", "Predictions are frequently skewed by recent single-match results, regional media popularity, or fan sentiment rather than objective data.")
    add_bullet("3. Unquantified Pressure & Clutch Factors: ", "Conventional statistics fail to isolate how teams perform specifically during high-pressure knockout stages (semi-finals and finals).")
    add_bullet("4. Unadjusted Host Venue Adaptability: ", "The 2027 World Cup will be held in South Africa, Zimbabwe, and Namibia. Pitch conditions in southern Africa favor specific pace, bounce, and seam metrics that require explicit venue adaptability weighting.")
    add_bullet("5. Lack of Human vs AI Benchmarking: ", "Fans have no structured platform to record their manual tournament predictions and benchmark their personal accuracy against machine learning algorithms over historical time horizons.")

    # -------------------------------------------------------------
    # 4. MARKET OPPORTUNITY
    # -------------------------------------------------------------
    add_h1("4. Market Opportunity & Target Audience")
    p = doc.add_paragraph()
    p.add_run("The platform addresses five distinct target user segments across the global sports analytics landscape:")

    add_bullet("• Cricket Fans & Enthusiasts: ", "Seeking interactive historical archives, clean visual stats, and a gamified Raj vs AI challenge to test their personal cricket knowledge.")
    add_bullet("• Sports Data Journalists & Content Creators: ", "Requiring empirical feature importance weights, team side-by-side matrices, and quantitative backtesting data to support editorial articles.")
    add_bullet("• Academic Computer Science Evaluators: ", "Looking for a rigorous, full-stack reference project implementing modern software architecture, REST APIs, and cross-validated ML estimators.")
    add_bullet("• Fantasy Sports Strategists: ", "Utilizing 3-year bilateral form ratings, batting depth metrics, and bowling efficiency indices to inform squad selection.")

    # -------------------------------------------------------------
    # 5. USER PERSONAS
    # -------------------------------------------------------------
    add_h1("5. User Personas")
    up_headers = ["Persona Profile", "Target User Role", "Primary Goals & Pain Points", "Key Platform Touchpoints"]
    up_widths = [1.5, 1.4, 2.2, 1.4]
    up_data = [
        ["Persona 1 — Cricket Fan (Arjun)", "Casual Sports Fan", "Wants quick predictions, historical winner lists, and fun prediction tracking without overwhelming math.", "Home Dashboard, History Archive, Raj vs AI page."],
        ["Persona 2 — Sports Analyst (Priya)", "Domain Journalist", "Needs quantitative team comparisons, knockout win rates, feature weights, and API endpoints.", "Team Analytics, AI Analysis view, FastAPI REST API."],
        ["Persona 3 — CS Student (Rohan)", "Data Science Student", "Wants transparent ML code, feature engineering formulas, cross-validation metrics, and clean pipeline.", "ML Pipeline, Feature Engineering, Documentation."],
        ["Persona 4 — Evaluator (Prof. Sharma)", "B.Tech Evaluator", "Evaluates system architecture, working functionality, zero runtime errors, and empirical accuracy.", "Full Stack UI, FastAPI Swagger Docs, Backtest metrics."]
    ]
    t_up = doc.add_table(rows=1, cols=4)
    style_table(t_up, up_widths, up_headers, up_data)

    # -------------------------------------------------------------
    # 6. COMPETITIVE ANALYSIS
    # -------------------------------------------------------------
    add_h1("6. Competitive Analysis")
    ca_headers = ["Feature Capability", "ESPNcricinfo / Cricbuzz", "Generic Sports Apps", "Our Platform (Raj vs AI)"]
    ca_widths = [1.8, 1.5, 1.5, 1.7]
    ca_data = [
        ["Historical World Cup Records", "Comprehensive raw scorecards", "Basic winner lists", "Structured 1983-2027 archive + drilldown modal"],
        ["Supervised ML Predictions", "Proprietary / Closed algorithms", "None or basic heuristics", "Transparent 10-feature Softmax ML classifier"],
        ["Raj vs AI Dual Engine", "None", "None", "Interactive 1983-2027 comparison + Scoreboard"],
        ["Feature Weight Transparency", "Hidden from public", "None", "Full 10-feature importance breakdown"],
        ["Side-by-Side Team Matrix", "Multi-page manual lookup", "Basic head-to-head text", "Side-by-side 9-metric matrix (2-4 teams)"],
        ["Offline Resilience", "Requires internet", "Requires internet", "Built-in fallback client dataset in api.js"],
        ["Academic Open Source Code", "Commercial closed source", "Commercial closed source", "100% open-source B.Tech reference project"]
    ]
    t_ca = doc.add_table(rows=1, cols=4)
    style_table(t_ca, ca_widths, ca_headers, ca_data)

    # -------------------------------------------------------------
    # 7. FUNCTIONAL REQUIREMENTS
    # -------------------------------------------------------------
    add_h1("7. Functional Requirements")
    fr_headers = ["ID", "Module Name", "Detailed Functional Specification", "Priority", "Status"]
    fr_widths = [0.8, 1.6, 2.7, 0.7, 0.7]
    fr_data = [
        ["FR-01", "Cricket Splash Screen", "Render fixed full-screen #000000 overlay with SVG bowling animation, stumps hit, and 4-step text for 2.5s on initial launch/refresh.", "High", "Implemented"],
        ["FR-02", "Home Overview Dashboard", "Display 2027 AI Champion card, Top 5 contenders leaderboard, tournament counters, and 1983-2023 champions honor roll.", "High", "Implemented"],
        ["FR-03", "World Cup History Archive", "Render 1983-2027 tournament table with search filter across years, teams, hosts, and match inspection modal.", "High", "Implemented"],
        ["FR-04", "Raj vs AI Prediction Engine", "Allow user to select predictions for 1983-2027, save in LocalStorage, compute live accuracy scoreboards and leader standings.", "High", "Implemented"],
        ["FR-05", "Team Analytics Comparison", "Compare 2 to 4 teams side-by-side across 9 quantitative historical and recent performance metrics.", "Medium", "Implemented"],
        ["FR-06", "AI Model Metrics View", "Display 5-fold CV metrics (Acc, Prec, Rec, F1), algorithm comparison table, 10 feature weights, and probabilities.", "High", "Implemented"],
        ["FR-07", "2027 Special Pick Highlight", "Display highlighted Raj's Pick vs AI's Pick card, agreement status badge, and 2027 team probability rankings.", "High", "Implemented"],
        ["FR-08", "Data Sources & Disclaimer", "Document official data citations (ICC, ESPNcricinfo) and present concise ML prediction disclaimer.", "Medium", "Implemented"],
        ["FR-09", "FastAPI REST API Server", "Expose 7 REST endpoints returning structured JSON data for predictions, history, teams, and metrics.", "High", "Implemented"],
        ["FR-10", "Offline Fallback Client", "Implement fallback dataset in api.js to guarantee 100% UI uptime even if backend server is offline.", "High", "Implemented"],
        ["FR-11", "Live Webhook Odds Updates", "Stream real-time match odds updates during live games via webhooks.", "Low", "Planned (V2)"],
        ["FR-12", "User Auth & Leaderboards", "User sign-in, JWT auth, and global multi-user Raj vs AI score ranking.", "Low", "Planned (V3)"]
    ]
    t_fr = doc.add_table(rows=1, cols=5)
    style_table(t_fr, fr_widths, fr_headers, fr_data)

    # -------------------------------------------------------------
    # 8. NON-FUNCTIONAL REQUIREMENTS
    # -------------------------------------------------------------
    add_h1("8. Non-Functional Requirements")
    add_bullet("1. Latency & Performance: ", "The React Vite production bundle compiles in 1.17 seconds (gzip JS size: 60.68 kB). REST API endpoints respond in <30ms on local ASGI server.")
    add_bullet("2. High Reliability & Fault Tolerance: ", "The frontend implements a dual-tier data fetcher in api.js. If backend API requests fail or time out, static fallback datasets render seamlessly, guaranteeing 100% UI uptime.")
    add_bullet("3. Responsive UX & Visual Accessibility: ", "Designed with a high-contrast dark theme (bg-slate-950, emerald #10b981, trophy gold #f59e0b) adhering to WCAG 2.1 AA guidelines. Responsive across mobile, tablet, and desktop screens.")
    add_bullet("4. Model Reproducibility: ", "Fixed random_state=42 across NumPy, KFold, and Scikit-learn estimators guarantees identical cross-validation metrics and predictions upon retraining.")
    add_bullet("5. Maintainability & Code Quality: ", "Structured React component hierarchy (components/, pages/, services/) and PEP 8 compliant Python code.")

    # -------------------------------------------------------------
    # 9. USER STORIES
    # -------------------------------------------------------------
    add_h1("9. User Stories")
    add_bullet("• User Story 1 (Cricket Fan): ", "As a cricket fan, I want to select my predicted winner for the 1999 World Cup, so that I can see if my choice beat the AI model's prediction.")
    add_bullet("  Acceptance Criteria: ", "Dropdown selector allows year 1999 selection -> Pick Australia -> Submit -> Result marked 'Correct' for both Raj and AI.")
    add_bullet("• User Story 2 (Sports Analyst): ", "As a sports analyst, I want to inspect feature importance weights, so that I can understand the mathematical rationale behind Australia's 31.3% win probability.")
    add_bullet("  Acceptance Criteria: ", "AI Analysis view renders 10 features with exact percentage weights (e.g., Knockout Stage Win Rate: 23.83%).")
    add_bullet("• User Story 3 (B.Tech Evaluator): ", "As a project evaluator, I want to view cross-validation metrics across all 4 ML algorithms, so that I can confirm model selection rigor.")
    add_bullet("  Acceptance Criteria: ", "Comparison table displays Accuracy, Precision, Recall, and F1-Score for Logistic Regression, Random Forest, Gradient Boosting, and Extra Trees.")

    # -------------------------------------------------------------
    # 10. USE CASES
    # -------------------------------------------------------------
    add_h1("10. Use Cases")
    p = doc.add_paragraph()
    p.add_run("Use Case UC-01: Compare Raj vs AI Predictions\n").bold = True
    add_bullet("• Primary Actor: ", "User (Raj)")
    add_bullet("• Preconditions: ", "App loaded in browser, LocalStorage accessible.")
    add_bullet("• Main Success Flow: ", "1. User navigates to Raj vs AI tab. 2. Selects year '2023' and pick 'India'. 3. Clicks 'Submit Prediction'. 4. System saves pick in LocalStorage, compares Raj's pick ('India') and AI's pick ('India') with Actual Winner ('Australia'). 5. Both marked Incorrect, scoreboard re-renders.")
    add_bullet("• Alternative Flow: ", "If user selects 2027, Actual Winner renders 'TBD', Raj & AI Results render 'Pending', and 2027 pick card renders 'Raj and AI agree'.")

    add_h2("Use Case Diagram Description")
    p_ucd = doc.add_paragraph()
    p_ucd.add_run("Actors (User / Analyst) initiate actions through the React UI boundary. Requests pass to Frontend Page Handlers (RajVsAi, History, TeamAnalytics, AiAnalysis), which fetch dataset JSON/CSV models via api.js from FastAPI Endpoints (/api/prediction/2027, /api/history, /api/teams). Endpoints query Scikit-Learn Model Binaries (best_model.pkl, scaler.pkl) to serve structured analytics.")

    # -------------------------------------------------------------
    # 11. SYSTEM WORKFLOWS
    # -------------------------------------------------------------
    add_h1("11. System Workflows")
    add_h2("1. User Navigation Workflow")
    add_code("Open Browser → Solid Black Splash Screen (2.5s) → Home Dashboard → Select Tab (History / Raj vs AI / Teams / AI Analysis) → Render View")

    add_h2("2. Machine Learning Training Workflow")
    add_code("Raw CSV Load → Normalize 10 Features → Generate Jitter Candidates (176 samples) → StandardScaler → 5-Fold Stratified CV → Select Logistic Regression → Refit Model → Softmax Temperature Scale → Save Artifacts (best_model.pkl, prediction_2027.json)")

    add_h2("3. Raj vs AI Dual Engine Workflow")
    add_code("Select Year (1983-2027) → Pick Team → Store in LocalStorage → Retrieve AI Prediction & Actual Winner → Calculate Accuracy Scoreboard → Update Leader Banner ('AI Leading' / 'Raj Leading' / 'Tie')")

    # -------------------------------------------------------------
    # 12. AI ARCHITECTURE
    # -------------------------------------------------------------
    add_h1("12. AI Architecture")
    p_ai = doc.add_paragraph()
    p_ai.add_run("The artificial intelligence architecture is structured into 7 distinct sequential processing stages:")

    add_code(
        "Raw Datasets (CSV/JSON)\n"
        "       │\n"
        "       ▼\n"
        "Feature Normalization (ml/feature_engineering.py)\n"
        "       │\n"
        "       ▼\n"
        "Synthetic Jitter Matrix Generation (176 samples)\n"
        "       │\n"
        "       ▼\n"
        "StandardScaler Transformation (models/scaler.pkl)\n"
        "       │\n"
        "       ▼\n"
        "Supervised ML Estimator (Logistic Regression)\n"
        "       │\n"
        "       ▼\n"
        "Softmax Temperature Calibration (T=3.5)\n"
        "       │\n"
        "       ▼\n"
        "FastAPI REST Endpoint (/api/prediction/2027) ──► React Frontend View"
    )

    add_h2("Mathematical Formulations")
    p_math = doc.add_paragraph()
    p_math.add_run("1. Feature Normalization Formula:\n").bold = True
    add_code("X_norm = (X - X_min) / (X_max - X_min)")
    p_math2 = doc.add_paragraph()
    p_math2.add_run("2. Softmax Temperature Probability Formula:\n").bold = True
    add_code("P(Team_i) = exp(z_i * T) / ∑ exp(z_j * T)   [where T = 3.5, z_i = log-odds]")

    # -------------------------------------------------------------
    # 13. SYSTEM ARCHITECTURE
    # -------------------------------------------------------------
    add_h1("13. System Architecture")
    p = doc.add_paragraph()
    p.add_run("The application follows a clean 4-tier microservice-inspired architecture:")
    add_bullet("1. Presentation Layer: ", "React 18 SPA built with Vite 5, Tailwind CSS 3, Lucide icons.")
    add_bullet("2. Application API Layer: ", "FastAPI 0.141 Python server running on Uvicorn ASGI on port 8000.")
    add_bullet("3. Machine Learning Layer: ", "Scikit-Learn 1.6 model engine (best_model.pkl, scaler.pkl).")
    add_bullet("4. Data Storage Layer: ", "JSON and CSV file datasets (world_cup_history.json, team_wc_performance.csv, recent_odi_stats.csv).")

    # -------------------------------------------------------------
    # 14. DATABASE & STORAGE DESIGN
    # -------------------------------------------------------------
    add_h1("14. Database & Storage Design")
    p_db = doc.add_paragraph()
    p_db.add_run("Current Implementation Note: ").bold = True
    p_db.add_run("The current project implementation does not use a traditional relational SQL database. Data is stored using structured project dataset files (world_cup_history.json, team_wc_performance.csv, recent_odi_stats.csv) and LocalStorage for client persistence.\n\n")
    p_db.add_run("Proposed Future Relational Schema (V2 Target):\n").bold = True
    add_bullet("• Teams Table: ", "team_id (PK, INT), team_name (VARCHAR), icc_rank (INT), recent_win_pct (FLOAT), batting_rating (FLOAT), bowling_rating (FLOAT).")
    add_bullet("• Tournaments Table: ", "tournament_id (PK, INT), year (INT), host_country (VARCHAR), winner_team_id (FK), runner_up_id (FK).")
    add_bullet("• UserPredictions Table: ", "pred_id (PK, INT), user_id (VARCHAR), tournament_id (FK), predicted_team_id (FK), created_at (TIMESTAMP).")

    # -------------------------------------------------------------
    # 15. API DESIGN & SPECIFICATION
    # -------------------------------------------------------------
    add_h1("15. API Design & Specification")
    api_headers = ["Endpoint URL", "Method", "Query Params", "Purpose & Response Description"]
    api_widths = [1.8, 0.8, 1.4, 2.5]
    api_data = [
        ["/", "GET", "None", "API root check. Returns status online, title, and docs link."],
        ["/api/health", "GET", "None", "Health check. Returns {'status': 'healthy'}."],
        ["/api/prediction/2027", "GET", "None", "Returns 2027 ML predictions, Softmax team probabilities, feature importances, explanation, and backtest array."],
        ["/api/history", "GET", "year (int, optional)", "Returns World Cup history records (1983-2027). Filters by year."],
        ["/api/teams", "GET", "None", "Returns merged historical performance and recent stats for all teams."],
        ["/api/teams/compare", "GET", "teams (str)", "Compares two or more teams side-by-side on 9 quantitative metrics."],
        ["/api/model/metrics", "GET", "None", "Returns validation metrics across all 4 trained ML algorithms and feature weights."]
    ]
    t_api = doc.add_table(rows=1, cols=4)
    style_table(t_api, api_widths, api_headers, api_data)

    # -------------------------------------------------------------
    # 16. DASHBOARD SPECIFICATIONS
    # -------------------------------------------------------------
    add_h1("16. Dashboard Specifications")
    add_bullet("1. Home Dashboard: ", "2027 Champion highlight card (Australia 31.3%), Top 5 contenders grid, historical stats counters.")
    add_bullet("2. World Cup History: ", "1983-2027 history table, real-time search filter input, inspection modal.")
    add_bullet("3. Raj vs AI Page: ", "Year & team dropdown selector, LocalStorage submit button, Scoreboard, Leader banner, 1983-2027 comparison table, 2027 pick card.")
    add_bullet("4. Team Analytics: ", "Multi-team selector pills, 9-metric side-by-side comparison table.")
    add_bullet("5. AI Analysis: ", "Model validation grid (Acc, Prec, Rec, F1), 4-algo comparison table, feature importance list, 10-team probability ranking.")
    add_bullet("6. 2027 Prediction View: ", "Raj's Pick vs AI's Pick card, agreement status badge, 2027 probability table.")
    add_bullet("7. Data Sources Page: ", "Citations list and statistical ML disclaimer.")

    # -------------------------------------------------------------
    # 17. WIREFRAME DESCRIPTIONS
    # -------------------------------------------------------------
    add_h1("17. UI Wireframe Layout Descriptions")
    add_code(
        "[ Top Navbar: Home | History | Raj vs AI | Team Analysis | AI Analysis | 2027 Prediction | Data Sources ]\n"
        "─────────────────────────────────────────────────────────────────────────────────────────────────────\n"
        "  [ Header Banner: 'Raj vs AI — ODI World Cup Prediction History (1983–2027)' ]\n"
        "  [ Form Card: Select Year [1983 ▼] | Select Team [India ▼] | ( Submit Prediction ) ]\n"
        "  [ Dynamic Scoreboard Card: Raj: 5/11 (45.5%) | AI: 8/11 (72.7%) | Leader: 'AI is currently leading!' ]\n"
        "  [ 2027 Special Card: RAJ'S PICK: India | AI'S PICK: Australia | Status: 'Raj and AI disagree' ]\n"
        "  [ Comparison Table: Year | Raj Pick | AI Pick | Actual Winner | Raj Result | AI Result ]\n"
        "─────────────────────────────────────────────────────────────────────────────────────────────────────\n"
        "[ Footer: Built with React, Vite, Tailwind CSS, Python FastAPI & Scikit-learn. Built by Raj Malviya ]"
    )

    # -------------------------------------------------------------
    # 18. AI MODELS
    # -------------------------------------------------------------
    add_h1("18. Machine Learning Model Analysis")
    ml_headers = ["Algorithm", "Accuracy", "Precision", "Recall", "F1-Score", "Hyperparameters", "Selection Status"]
    ml_widths = [1.5, 0.8, 0.8, 0.8, 0.8, 1.2, 0.6]
    ml_data = [
        ["Logistic Regression", "74.38%", "66.46%", "46.27%", "0.5070", "C=1.0, random_state=42", "SELECTED BEST"],
        ["Gradient Boosting", "71.88%", "54.72%", "46.66%", "0.4982", "n_est=100, lr=0.08, depth=3", "Evaluated"],
        ["Random Forest", "70.63%", "52.97%", "44.22%", "0.4552", "n_est=100, depth=5", "Evaluated"],
        ["Extra Trees", "71.88%", "62.86%", "35.03%", "0.3992", "n_est=100, depth=5", "Evaluated"]
    ]
    t_ml = doc.add_table(rows=1, cols=7)
    style_table(t_ml, ml_widths, ml_headers, ml_data)

    p_reason = doc.add_paragraph()
    p_reason.paragraph_format.space_before = Pt(6)
    p_reason.add_run("Why Logistic Regression was Selected: ").bold = True
    p_reason.add_run("Logistic Regression achieved the highest cross-validated F1-score (0.5070) and validation accuracy (74.38%). Its linear log-odds decision boundary provides smooth, well-calibrated probability outputs when transformed via Softmax, preventing over-confident extreme predictions.")

    # -------------------------------------------------------------
    # 19. SEARCH ALGORITHMS
    # -------------------------------------------------------------
    add_h1("19. Search Algorithms & Data Filtering")
    p_src = doc.add_paragraph()
    p_src.add_run("Implementation Note: ").bold = True
    p_src.add_run("No dedicated graph or text search algorithm (such as A*, Dijkstra, or Elasticsearch) is currently implemented because the application's primary functionality is statistical analysis and machine learning prediction. Real-time text filtering on the World Cup History page is executed using client-side JavaScript String.prototype.includes() matching across years, teams, and host countries.")

    # -------------------------------------------------------------
    # 20. KNOWLEDGE REPRESENTATION
    # -------------------------------------------------------------
    add_h1("20. Knowledge Representation")
    p_kn = doc.add_paragraph()
    p_kn.add_run("Domain knowledge is represented as structured feature vectors (10 normalized metrics between 0.0 and 1.0) and JSON tournament models encoding historical outcomes, scorelines, top scorers, and backtested predictions.")

    # -------------------------------------------------------------
    # 21. MACHINE LEARNING PIPELINE
    # -------------------------------------------------------------
    add_h1("21. Machine Learning Pipeline")
    add_code("Raw Data (CSV) → Feature Normalization → Synthetic Jitter Dataset (176 samples) → StandardScaler → 5-Fold Stratified CV → Logistic Regression Refit → Softmax Temperature Scaling → JSON Export")

    # -------------------------------------------------------------
    # 22. ANALYTICS MODULE
    # -------------------------------------------------------------
    add_h1("22. Analytics Module")
    add_bullet("• Backtest Accuracy Analytics: ", "Calculates AI historical backtest accuracy across past tournaments (72.7% accuracy, 8/11 correct).")
    add_bullet("• Head-to-Head Raj vs AI Analytics: ", "Real-time computation of Raj's accuracy vs AI accuracy.")
    add_bullet("• Team Comparison Analytics: ", "Side-by-side metric normalization across batting, bowling, knockout clutch, and African venue ratings.")

    # -------------------------------------------------------------
    # 23. SECURITY
    # -------------------------------------------------------------
    add_h1("23. Security Posture")
    add_bullet("• Input Validation: ", "FastAPI query parameter typing (e.g., year: Optional[int]) prevents invalid payload types.")
    add_bullet("• CORS Isolation: ", "CORSMiddleware configured for API protection.")
    add_bullet("• Current Limitation Note: ", "No user authentication or JWT security is currently implemented. The system operates as an open academic analytics dashboard.")

    # -------------------------------------------------------------
    # 24. TESTING STRATEGY
    # -------------------------------------------------------------
    add_h1("24. Testing Strategy & Case Matrix")
    tc_headers = ["Test ID", "Feature Tested", "Input Scenario", "Expected Result", "Actual Result", "Status"]
    tc_widths = [0.8, 1.2, 1.5, 1.2, 1.2, 0.6]
    tc_data = [
        ["TC-01", "Splash Screen", "Initial page load / refresh", "Full black overlay for 2.5s", "Rendered #000000 overlay 2.5s", "PASS"],
        ["TC-02", "Raj Prediction", "Select 2023, Pick 'India'", "Saved in LocalStorage, score update", "Saved & updated scoreboards", "PASS"],
        ["TC-03", "2027 AI API", "GET /api/prediction/2027", "Return Australia 31.3%, top 5 array", "Returned status 200 JSON payload", "PASS"],
        ["TC-04", "Team Compare", "GET /api/teams/compare?teams=India,Australia", "Return 2 team records", "Returned merged CSV metrics", "PASS"],
        ["TC-05", "Vite Build", "npm run build", "Production bundle compilation", "Built in 1.17s with 0 errors", "PASS"]
    ]
    t_tc = doc.add_table(rows=1, cols=6)
    style_table(t_tc, tc_widths, tc_headers, tc_data)

    # -------------------------------------------------------------
    # 25. KPIS
    # -------------------------------------------------------------
    add_h1("25. Key Performance Indicators (KPIs)")
    add_bullet("• Technical KPIs: ", "5-Fold CV Accuracy (74.38%), API Response Time (<30ms), Frontend Build Time (1.17s), JS Bundle Size (212 kB).")
    add_bullet("• Academic KPIs: ", "Historical Backtesting Score (72.7%), Data Completeness (100% across 12 World Cups).")

    # -------------------------------------------------------------
    # 26. SUCCESS METRICS
    # -------------------------------------------------------------
    add_h1("26. Success Metrics")
    add_bullet("1. Zero Runtime Errors: ", "Clean console logs and error-free execution across all 7 frontend page views.")
    add_bullet("2. State Persistence: ", "100% reliable LocalStorage client prediction saving for Raj's picks.")
    add_bullet("3. Accurate Data Rendering: ", "Factual alignment with official ICC World Cup archives.")

    # -------------------------------------------------------------
    # 27. ROADMAP
    # -------------------------------------------------------------
    add_h1("27. Development Roadmap")
    rm_headers = ["Phase", "Milestones & Scope", "Target Deliverables", "Status"]
    rm_widths = [0.8, 3.2, 1.8, 0.7]
    rm_data = [
        ["MVP", "1983-2027 history dataset, 10-feature ML model, Raj vs AI engine, FastAPI, React dark UI.", "Fully functional local web app", "COMPLETED"],
        ["V1", "Improved feature weights, expanded team stats, exportable PDF report.", "Enhanced analytical dashboard", "Current"],
        ["V2", "Live ICC webhook API integration, ball-by-ball win probability graph, relational Postgres DB.", "Real-time sports data portal", "Planned"],
        ["V3", "10,000 Monte Carlo match simulations, mobile React Native app, multi-user prediction challenges.", "Enterprise sports AI platform", "Planned"]
    ]
    t_rm = doc.add_table(rows=1, cols=4)
    style_table(t_rm, rm_widths, rm_headers, rm_data)

    # -------------------------------------------------------------
    # 28. RISKS & MITIGATION
    # -------------------------------------------------------------
    add_h1("28. Risk & Mitigation Matrix")
    rk_headers = ["Risk Description", "Impact", "Probability", "Mitigation Strategy"]
    rk_widths = [2.0, 0.8, 0.8, 2.9]
    rk_data = [
        ["Small Historical Sample Size (11 WCs)", "High", "High", "Generated synthetic jitter training candidates to train robust classifiers."],
        ["Knockout Match Non-Determinism", "Medium", "High", "Calculated Softmax probability distributions rather than hard binary outputs."],
        ["Backend Service Offline", "High", "Low", "Built offline fallback dataset in frontend api.js to ensure 100% UI uptime."]
    ]
    t_rk = doc.add_table(rows=1, cols=4)
    style_table(t_rk, rk_widths, rk_headers, rk_data)

    # -------------------------------------------------------------
    # 29. FUTURE ENHANCEMENTS
    # -------------------------------------------------------------
    add_h1("29. Future Enhancements")
    add_bullet("1. Real-Time Webhook Integration: ", "Connecting live match score APIs to update win probabilities ball-by-ball.")
    add_bullet("2. Player-Level Elo Ratings: ", "Incorporating individual batting/bowling Elo ratings and squad injury status.")
    add_bullet("3. Monte Carlo Tournament Simulator: ", "Simulating 10,000 automated tournament match iterations.")

    # -------------------------------------------------------------
    # 30. APPENDIX
    # -------------------------------------------------------------
    add_h1("30. Appendix")
    
    add_h2("A. Technology Stack Reference")
    p_app_a = doc.add_paragraph()
    p_app_a.add_run("React 18.2, Vite 5.1, Tailwind CSS 3.4, Python 3.10+, FastAPI 0.141, Scikit-Learn 1.6, Pandas 2.2, NumPy 2.1.")

    add_h2("B. Complete Project Directory Tree")
    p_app_b = doc.add_paragraph()
    r_tree = p_app_b.add_run(
        "cricket-analytics-2027/\n"
        "├── data/ (world_cup_history.json, team_wc_performance.csv, recent_odi_stats.csv)\n"
        "├── ml/ (feature_engineering.py, train_model.py)\n"
        "├── models/ (best_model.pkl, scaler.pkl, prediction_2027.json)\n"
        "├── backend/ (main.py, requirements.txt)\n"
        "├── frontend/ (index.html, vite.config.js, package.json, src/App.jsx, src/components/, src/pages/)\n"
        "└── README.md"
    )
    r_tree.font.name = 'Consolas'
    r_tree.font.size = Pt(8.5)

    add_h2("C. Dataset Schema Breakdown")
    p_app_c = doc.add_paragraph()
    p_app_c.add_run("team_wc_performance.csv: Team, Titles, Finals_Appearances, Semi_Finals, Matches_Played, Matches_Won, Matches_Lost, Win_Percentage, Batting_Avg, Bowling_Avg, Net_Run_Rate, Knockout_Win_Pct.")

    add_h2("D. Model Features")
    p_app_d = doc.add_paragraph()
    p_app_d.add_run("10 Features: feat_wc_win_pct, feat_title_exp, feat_knockout_clutch, feat_recent_win_pct, feat_batting_strength, feat_bowling_strength, feat_consistency, feat_host_adaptability, feat_h2h_top5, feat_power_index.")

    add_h2("E. API Reference")
    p_app_e = doc.add_paragraph()
    p_app_e.add_run("GET / | GET /api/health | GET /api/prediction/2027 | GET /api/history | GET /api/teams | GET /api/teams/compare | GET /api/model/metrics.")

    add_h2("F. UI Screenshots Placeholders")
    p_app_f = doc.add_paragraph()
    p_app_f.add_run("[Placeholder: Solid Black Cricket Splash Screen (2.5s)]\n[Placeholder: Home Dashboard & 2027 Winner Card]\n[Placeholder: Raj vs AI Prediction Matrix & Scoreboard]\n[Placeholder: Team Side-by-Side Comparison Tool]\n[Placeholder: AI Analysis & Feature Importance Bar Chart]")

    add_h2("G. Technical Glossary")
    p_app_g = doc.add_paragraph()
    p_app_g.add_run("• ODI: One Day International cricket match (50 overs per side).\n• Softmax: Mathematical function that normalizes a vector of K real numbers into a probability distribution summing to 1.0.\n• Stratified K-Fold CV: Cross-validation technique that preserves target class proportions across all folds.\n• FastAPI: Modern, high-performance web framework for building APIs with Python.")

    # Save Word documents
    path1 = "/Users/rajmalviya/Desktop/cricket-analytics-2027/PROJECT_DOCUMENTATION.docx"
    path2 = "/Users/rajmalviya/Desktop/cricket-analytics-2027/ODI_World_Cup_2027_Academic_Documentation.docx"
    doc.save(path1)
    doc.save(path2)
    print(f"Successfully created expanded Word documentation at:\n  - {path1}\n  - {path2}")

if __name__ == "__main__":
    create_document()
