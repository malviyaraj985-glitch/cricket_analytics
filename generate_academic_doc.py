import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_document():
    doc = Document()
    
    # Page setup - Margins (1 inch all sides)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Color Palette - Professional Academic Slate & Emerald
    PRIMARY_COLOR = RGBColor(16, 185, 129)    # Emerald Green #10B981
    SECONDARY_COLOR = RGBColor(245, 158, 11)  # Trophy Gold #F59E0B
    DARK_BG = RGBColor(15, 23, 42)            # Slate Dark #0F172A
    TEXT_COLOR = RGBColor(30, 41, 59)          # Slate Text #1E293B
    MUTED_COLOR = RGBColor(100, 116, 139)      # Slate Muted #64748B

    # Helper function for setting cell shading
    def set_cell_background(cell, fill_hex):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # Helper function for setting cell margins/padding
    def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
            node = OxmlElement(margin)
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    # Set Default Font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = TEXT_COLOR

    # Function to add styled headings
    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = SECONDARY_COLOR
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = DARK_BG
        return p

    def add_bullet(bold_prefix="", text=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(3)
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.bold = True
        p.add_run(text)
        return p

    def style_table(table, col_widths, headers, data):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        for i, header_text in enumerate(headers):
            hdr_cells[i].text = header_text
            set_cell_background(hdr_cells[i], "0F172A")
            set_cell_margins(hdr_cells[i], top=120, bottom=120, left=140, right=140)
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
        
        for row_idx, row_data in enumerate(data):
            row_cells = table.add_row().cells
            bg_color = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
            for i, val in enumerate(row_data):
                row_cells[i].text = str(val)
                set_cell_background(row_cells[i], bg_color)
                set_cell_margins(row_cells[i], top=100, bottom=100, left=140, right=140)
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(9.5)
                        run.font.color.rgb = TEXT_COLOR
                        
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    # -------------------------------------------------------------
    # TITLE BLOCK
    # -------------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(10)
    title_p.paragraph_format.space_after = Pt(2)
    t_run = title_p.add_run("ODI World Cup 2027 AI Analytics & Raj vs AI Prediction Platform")
    t_run.font.size = Pt(24)
    t_run.font.bold = True
    t_run.font.color.rgb = PRIMARY_COLOR

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(8)
    s_run = sub_p.add_run("AI/ML-Based ODI World Cup Historical Analysis and Winner Prediction System")
    s_run.font.size = Pt(14)
    s_run.font.bold = True
    s_run.font.color.rgb = SECONDARY_COLOR

    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_after = Pt(18)
    m_run = meta_p.add_run("Author: Raj Malviya  |  Technology: React + Vite + Tailwind CSS + Python + FastAPI + Scikit-learn")
    m_run.font.size = Pt(10)
    m_run.font.italic = True
    m_run.font.color.rgb = MUTED_COLOR

    # -------------------------------------------------------------
    # 1. EXECUTIVE SUMMARY
    # -------------------------------------------------------------
    add_h1("1. Executive Summary")
    p = doc.add_paragraph()
    p.add_run("The ")
    p.add_run("ODI World Cup 2027 AI Analytics & Raj vs AI Prediction Platform").bold = True
    p.add_run(" is a full-stack data analytics and machine learning application developed to evaluate 40 years of ICC Men's ODI World Cup historical records (1983–2023), team strength indices, and regional venue adaptability. The system solves the problem of subjective sports commentary by implementing a quantitative ML pipeline using 5-fold Stratified Cross-Validation across 4 supervised algorithms. Selected model ")
    p.add_run("Logistic Regression").bold = True
    p.add_run(" (Validation Accuracy: 74.38%, Backtest Accuracy: 72.7%) computes Softmax win probability distributions for the 2027 World Cup in South Africa, Zimbabwe, and Namibia (predicting Australia at 31.3% and India at 23.9%). The core feature, ")
    p.add_run("Raj vs AI").bold = True
    p.add_run(", enables interactive head-to-head prediction tracking between user choices (Raj) and machine learning predictions across all 12 World Cups (1983–2027) with LocalStorage client persistence.")

    # -------------------------------------------------------------
    # 2. VISION
    # -------------------------------------------------------------
    add_h1("2. Vision")
    p = doc.add_paragraph()
    p.add_run("The platform envisions bridging the gap between empirical sports data science and fan engagement. By replacing bias with cross-validated machine learning features, the long-term vision encompasses real-time tournament simulations, ball-by-ball predictive updates, and transparent model explainability for sports analysts, fans, and researchers.")

    # -------------------------------------------------------------
    # 3. BUSINESS PROBLEM
    # -------------------------------------------------------------
    add_h1("3. Business Problem")
    p = doc.add_paragraph()
    p.add_run("Traditional sports predictions rely heavily on speculative commentary, fan bias, and unstructured historical data. Key challenges addressed include:")
    add_bullet("1. Manual Complexity: ", "Analyzing 556 historical World Cup matches across 11 editions is difficult without structured feature vectors.")
    add_bullet("2. Unquantified Pressure Factors: ", "Evaluating knockout clutch conversion rates under high pressure.")
    add_bullet("3. Lack of Human vs AI Benchmarking: ", "No existing simple platform allows fans to systematically benchmark human predictions against machine learning models over 40 years of history.")

    # -------------------------------------------------------------
    # 4. MARKET OPPORTUNITY
    # -------------------------------------------------------------
    add_h1("4. Market Opportunity")
    p = doc.add_paragraph()
    p.add_run("The platform serves global cricket analytics audiences, including:")
    add_bullet("• Cricket Fans & Enthusiasts: ", "Interactive tournament history inspection and gamified Raj vs AI prediction challenge.")
    add_bullet("• Sports Data Analysts & Journalists: ", "Quantitative team comparisons and ML feature importance weights.")
    add_bullet("• Academic Researchers & Students: ", "A complete reference architecture combining React 18, FastAPI, and Scikit-learn.")

    # -------------------------------------------------------------
    # 5. USER PERSONAS
    # -------------------------------------------------------------
    add_h1("5. User Personas")
    up_headers = ["Persona Name", "Role / Profile", "Primary Goals & Needs", "Platform Interaction Points"]
    up_widths = [1.5, 1.5, 2.0, 1.5]
    up_data = [
        ["Persona 1 — Cricket Fan", "General Sports Fan", "Wants simple predictions, historical winner lists, and fun prediction tracking.", "Home Dashboard, History Archive, Raj vs AI page."],
        ["Persona 2 — Sports Analyst", "Domain Expert", "Requires team stats, H2H records, feature importances, and model metrics.", "Team Analytics, AI Analysis page, REST API."],
        ["Persona 3 — Academic Student", "Data Science Student", "Wants transparent ML pipeline, datasets, cross-validation metrics, and clean code.", "ML Pipeline, Feature Engineering, Documentation."],
        ["Persona 4 — Project Evaluator", "B.Tech Evaluator", "Evaluates system architecture, working functionality, code quality, and accuracy.", "Full Stack UI, FastAPI Swagger Docs, Backtest metrics."]
    ]
    t_up = doc.add_table(rows=1, cols=4)
    style_table(t_up, up_widths, up_headers, up_data)

    # -------------------------------------------------------------
    # 6. COMPETITIVE ANALYSIS
    # -------------------------------------------------------------
    add_h1("6. Competitive Analysis")
    ca_headers = ["Feature Metric", "ESPNcricinfo / Cricbuzz", "Generic Sports Apps", "Our Platform (Raj vs AI)"]
    ca_widths = [1.8, 1.6, 1.5, 1.6]
    ca_data = [
        ["Historical WC Records", "Comprehensive raw data", "Basic score tables", "Structured 1983-2027 archive + drilldown"],
        ["Machine Learning Predictions", "Proprietary/Limited", "None or heuristic", "Transparent 10-feature Softmax ML model"],
        ["Human vs AI (Raj vs AI)", "None", "None", "Interactive 1983-2027 comparison + Scoreboard"],
        ["ML Explainability Weights", "Hidden / Not provided", "None", "Full feature importance breakdown"],
        ["Team Side-by-Side Tool", "Multi-click stats lookup", "Basic head-to-head", "Side-by-side 9-metric matrix (2-4 teams)"],
        ["Academic Transparency", "Commercial / Closed", "Commercial", "100% open source B.Tech architecture"]
    ]
    t_ca = doc.add_table(rows=1, cols=4)
    style_table(t_ca, ca_widths, ca_headers, ca_data)

    # -------------------------------------------------------------
    # 7. FUNCTIONAL REQUIREMENTS
    # -------------------------------------------------------------
    add_h1("7. Functional Requirements")
    fr_headers = ["ID", "Requirement Description", "Detailed Specification", "Priority", "Status"]
    fr_widths = [0.8, 1.8, 2.5, 0.7, 0.7]
    fr_data = [
        ["FR-01", "Solid Black Cricket Splash Screen", "Render fixed full-screen #000000 splash screen with SVG bowling animation and text for 2.5s.", "High", "Implemented"],
        ["FR-02", "World Cup History Exploration", "Display 1983-2027 tournament table, search filter, and match inspection modal.", "High", "Implemented"],
        ["FR-03", "Raj vs AI Prediction System", "Allow user to pick predicted winners for 1983-2027, save in LocalStorage, compute scoreboards.", "High", "Implemented"],
        ["FR-04", "Team Analytics Comparison", "Compare 2-4 teams side-by-side on 9 quantitative historical & recent metrics.", "Medium", "Implemented"],
        ["FR-05", "AI Model Metrics & Features", "Display 5-fold CV metrics, algorithm comparison table, 10 feature weights, and probabilities.", "High", "Implemented"],
        ["FR-06", "2027 Prediction Highlight View", "Render Raj's Pick vs AI's Pick comparison card and full 10-team win probabilities.", "High", "Implemented"],
        ["FR-07", "FastAPI REST API Services", "Expose 7 REST endpoints returning JSON dataset records, predictions, and metrics.", "High", "Implemented"],
        ["FR-08", "Live Ball-by-Ball Tracking", "Update win probabilities during ongoing live matches.", "Low", "Planned (V2)"],
        ["FR-09", "User Auth & Global Leaderboards", "User sign-in and global multi-user Raj vs AI score ranking.", "Low", "Planned (V3)"]
    ]
    t_fr = doc.add_table(rows=1, cols=5)
    style_table(t_fr, fr_widths, fr_headers, fr_data)

    # -------------------------------------------------------------
    # 8. NON-FUNCTIONAL REQUIREMENTS
    # -------------------------------------------------------------
    add_h1("8. Non-Functional Requirements")
    add_bullet("1. Performance: ", "Frontend Vite production bundle builds in 1.17s (212 kB JS). API responses serve in under 30ms locally.")
    add_bullet("2. Reliability & Resilience: ", "Frontend api.js client includes built-in fallback datasets to ensure 100% UI uptime even if backend server is offline.")
    add_bullet("3. Usability: ", "Clean dark theme B.Tech aesthetic (bg-slate-950, emerald #10b981) responsive across desktop and mobile viewports.")
    add_bullet("4. Model Reproducibility: ", "Fixed random_state=42 across NumPy, KFold, and Scikit-learn estimators ensures deterministic training outputs.")

    # -------------------------------------------------------------
    # 9. USER STORIES
    # -------------------------------------------------------------
    add_h1("9. User Stories")
    add_bullet("• As a Cricket Fan: ", "I want to select my predicted winner for the 1999 World Cup, so that I can see if my pick beat the AI model.")
    add_bullet("• As a Sports Analyst: ", "I want to inspect the feature importance weights, so that I can understand why the AI predicts Australia as the top 2027 contender.")
    add_bullet("• As a Student/Evaluator: ", "I want to view the cross-validation metrics across all 4 ML algorithms, so that I can verify model selection rigor.")

    # -------------------------------------------------------------
    # 10. USE CASES
    # -------------------------------------------------------------
    add_h1("10. Use Cases")
    p = doc.add_paragraph()
    p.add_run("Use Case UC-01: Compare Raj vs AI Predictions\n").bold = True
    add_bullet("• Primary Actor: ", "User (Raj)")
    add_bullet("• Preconditions: ", "App loaded, LocalStorage accessible.")
    add_bullet("• Main Flow: ", "1. User navigates to Raj vs AI tab. 2. Selects year '2023' and pick 'India'. 3. Clicks 'Submit Prediction'. 4. System saves pick, compares with AI pick ('India') and Actual Winner ('Australia'). 5. Both marked Incorrect, scoreboard updates.")
    add_bullet("• Postconditions: ", "State stored in LocalStorage, scoreboard re-rendered.")

    # -------------------------------------------------------------
    # 11. WORKFLOWS
    # -------------------------------------------------------------
    add_h1("11. System Workflows")
    p_wf = doc.add_paragraph()
    p_wf.add_run("Raj vs AI Workflow:\n").bold = True
    p_wf.add_run("Select World Cup Year (1983-2027)  →  Raj Selects Team  →  Save in LocalStorage  →  Retrieve AI Prediction & Actual Winner  →  Compute Correct/Incorrect Result  →  Update Scoreboard & Leader Banner")

    # -------------------------------------------------------------
    # 12. AI ARCHITECTURE
    # -------------------------------------------------------------
    add_h1("12. AI Architecture")
    p_ai = doc.add_paragraph()
    p_ai.add_run("Data Layer (CSV/JSON)  →  Feature Engineering (10 Vectors)  →  StandardScaler  →  Supervised Classifier (Logistic Regression)  →  Softmax Temperature Scaling  →  FastAPI Endpoint (/api/prediction/2027)  →  React Frontend Display")

    # -------------------------------------------------------------
    # 13. SYSTEM ARCHITECTURE
    # -------------------------------------------------------------
    add_h1("13. System Architecture")
    add_bullet("• Presentation Layer: ", "React 18 SPA built with Vite 5, Tailwind CSS 3, Lucide icons.")
    add_bullet("• Application API Layer: ", "FastAPI 0.141 Python server running on Uvicorn ASGI on port 8000.")
    add_bullet("• Machine Learning Layer: ", "Scikit-Learn 1.6 model engine (best_model.pkl, scaler.pkl).")
    add_bullet("• Data Storage Layer: ", "JSON and CSV file datasets (world_cup_history.json, team_wc_performance.csv, recent_odi_stats.csv).")

    # -------------------------------------------------------------
    # 14. DATABASE DESIGN
    # -------------------------------------------------------------
    add_h1("14. Database & Storage Design")
    p_db = doc.add_paragraph()
    p_db.add_run("Current Implementation Note: ").bold = True
    p_db.add_run("The current project implementation does not use a traditional relational SQL database. Data is stored using structured project dataset files (world_cup_history.json, team_wc_performance.csv, recent_odi_stats.csv) and LocalStorage for client persistence.\n\n")
    p_db.add_run("Proposed Future Relational Schema (V2 Target):\n").bold = True
    add_bullet("• Teams Table: ", "team_id (PK), team_name, icc_rank, recent_win_pct, batting_rating, bowling_rating.")
    add_bullet("• Tournaments Table: ", "tournament_id (PK), year, host_country, winner_team_id (FK), runner_up_id (FK).")
    add_bullet("• UserPredictions Table: ", "pred_id (PK), user_id, tournament_id (FK), predicted_team_id (FK), created_at.")

    # -------------------------------------------------------------
    # 15. API DESIGN
    # -------------------------------------------------------------
    add_h1("15. API Design & Specification")
    api_headers = ["Endpoint", "Method", "Query Params", "Response Summary"]
    api_widths = [1.8, 0.8, 1.5, 2.4]
    api_data = [
        ["/", "GET", "None", "Status online, app title, docs link."],
        ["/api/health", "GET", "None", "Returns {'status': 'healthy'}."],
        ["/api/prediction/2027", "GET", "None", "Returns 2027 ML predictions, Softmax team probabilities, feature importances, explanation, and backtest array."],
        ["/api/history", "GET", "year (int, optional)", "Returns World Cup history records (1983-2027). Filters by year."],
        ["/api/teams", "GET", "None", "Returns merged historical performance and recent stats."],
        ["/api/teams/compare", "GET", "teams (str)", "Compares two or more teams side-by-side."],
        ["/api/model/metrics", "GET", "None", "Returns evaluation scores for all 4 trained ML algorithms and feature weights."]
    ]
    t_api = doc.add_table(rows=1, cols=4)
    style_table(t_api, api_widths, api_headers, api_data)

    # -------------------------------------------------------------
    # 16. DASHBOARD SPECIFICATIONS
    # -------------------------------------------------------------
    add_h1("16. Dashboard & Interface Specifications")
    add_bullet("1. Home Dashboard: ", "2027 Champion highlight card, Top 5 contenders grid, historical stats counters.")
    add_bullet("2. World Cup History: ", "1983-2027 history table, search filter input, inspection modal.")
    add_bullet("3. Raj vs AI Page: ", "Year & team dropdown selector, LocalStorage submit button, Scoreboard, Leader banner, 1983-2027 comparison table, 2027 pick card.")
    add_bullet("4. Team Analytics: ", "Multi-team selector pills, 9-metric side-by-side comparison table.")
    add_bullet("5. AI Analysis: ", "Model validation grid (Acc, Prec, Rec, F1), 4-algo comparison table, feature importance list, 10-team probability ranking.")
    add_bullet("6. 2027 Prediction View: ", "Raj's Pick vs AI's Pick card, agreement status badge, 2027 probability table.")
    add_bullet("7. Data Sources Page: ", "Citations list and statistical ML disclaimer.")

    # -------------------------------------------------------------
    # 17. WIREFRAME DESCRIPTIONS
    # -------------------------------------------------------------
    add_h1("17. UI Wireframe Layout Descriptions")
    p_wf = doc.add_paragraph()
    p_wf.add_run("Raj vs AI Wireframe Structure:\n").bold = True
    p_wf.add_run("[Top Navbar (7 Tabs)]\n  └── [Header Banner: 'Raj vs AI — ODI World Cup Prediction History']\n  └── [Prediction Form: 'Select Year ▼' | 'Select Team ▼' | (Submit Prediction)]\n  └── [Dynamic Scoreboard Card: Raj Score vs AI Score | Leader Banner]\n  └── [2027 Highlight Card: RAJ'S PICK vs AI'S PICK | Agree/Disagree Badge]\n  └── [1983–2027 Table: Year | Raj Pick | AI Pick | Actual Winner | Raj Result | AI Result]\n[Footer: 'Built by Raj Malviya']")

    # -------------------------------------------------------------
    # 18. AI MODELS
    # -------------------------------------------------------------
    add_h1("18. Machine Learning Model Analysis")
    ml_headers = ["Algorithm", "Accuracy", "Precision", "Recall", "F1-Score", "Hyperparameters", "Status"]
    ml_widths = [1.5, 0.8, 0.8, 0.8, 0.8, 1.2, 0.6]
    ml_data = [
        ["Logistic Regression", "74.38%", "66.46%", "46.27%", "0.5070", "C=1.0, random_state=42", "BEST"],
        ["Gradient Boosting", "71.88%", "54.72%", "46.66%", "0.4982", "n_est=100, lr=0.08, depth=3", "Eval"],
        ["Random Forest", "70.63%", "52.97%", "44.22%", "0.4552", "n_est=100, depth=5", "Eval"],
        ["Extra Trees", "71.88%", "62.86%", "35.03%", "0.3992", "n_est=100, depth=5", "Eval"]
    ]
    t_ml = doc.add_table(rows=1, cols=7)
    style_table(t_ml, ml_widths, ml_headers, ml_data)

    # -------------------------------------------------------------
    # 19. SEARCH ALGORITHMS
    # -------------------------------------------------------------
    add_h1("19. Search Algorithms & Data Filtering")
    p_src = doc.add_paragraph()
    p_src.add_run("Implementation Note: ").bold = True
    p_src.add_run("No dedicated graph or text search algorithm (such as A*, Dijkstra, or Elasticsearch) is currently implemented because the application's primary functionality is statistical analysis and machine learning prediction. Real-time text filtering on the World Cup History page is executed using client-side JavaScript String.prototype.includes() matching across years, teams, and host countries.")

    # -------------------------------------------------------------
    # 20. KNOWLEDGE REPRESENTATION
    # -------------------------------------------------------------
    add_h1("20. Knowledge Representation")
    p_kn = doc.add_paragraph()
    p_kn.add_run("Domain knowledge is represented as structured feature vectors (10 normalized metrics between 0.0 and 1.0) and JSON tournament models encoding historical outcomes, scorelines, top scorers, and backtested predictions.")

    # -------------------------------------------------------------
    # 21. MACHINE LEARNING PIPELINE
    # -------------------------------------------------------------
    add_h1("21. Machine Learning Pipeline")
    p_pipe = doc.add_paragraph()
    p_pipe.add_run("Raw Data (CSV)  →  Feature Normalization  →  Synthetic Jitter Dataset (176 samples)  →  StandardScaler  →  5-Fold Stratified CV  →  Logistic Regression Refit  →  Softmax Temperature Scaling  →  JSON Export")

    # -------------------------------------------------------------
    # 22. ANALYTICS MODULE
    # -------------------------------------------------------------
    add_h1("22. Analytics Module")
    add_bullet("• Backtest Accuracy Analytics: ", "Calculates AI historical backtest accuracy across past tournaments (72.7% accuracy, 8/11 correct).")
    add_bullet("• Head-to-Head Raj vs AI Analytics: ", "Real-time computation of Raj's accuracy vs AI accuracy.")
    add_bullet("• Team Comparison Analytics: ", "Side-by-side metric normalization across batting, bowling, knockout clutch, and African venue ratings.")

    # -------------------------------------------------------------
    # 23. SECURITY
    # -------------------------------------------------------------
    add_h1("23. Security Posture")
    add_bullet("• Input Validation: ", "FastAPI query parameter typing (e.g., year: Optional[int]) prevents invalid payload types.")
    add_bullet("• CORS Isolation: ", "CORSMiddleware configured for API protection.")
    add_bullet("• Current Limitation Note: ", "No user authentication or JWT security is currently implemented. The system operates as an open academic analytics dashboard.")

    # -------------------------------------------------------------
    # 24. TESTING STRATEGY
    # -------------------------------------------------------------
    add_h1("24. Testing Strategy & Case Matrix")
    tc_headers = ["Test ID", "Feature Tested", "Input Scenario", "Expected Result", "Actual Result", "Status"]
    tc_widths = [0.8, 1.2, 1.5, 1.2, 1.2, 0.6]
    tc_data = [
        ["TC-01", "Splash Screen", "Initial page load / refresh", "Full black overlay for 2.5s", "Rendered #000000 overlay 2.5s", "PASS"],
        ["TC-02", "Raj Prediction", "Select 2023, Pick 'India'", "Saved in LocalStorage, score update", "Saved & updated scoreboards", "PASS"],
        ["TC-03", "2027 AI API", "GET /api/prediction/2027", "Return Australia 31.3%, top 5 array", "Returned status 200 JSON payload", "PASS"],
        ["TC-04", "Team Compare", "GET /api/teams/compare?teams=India,Australia", "Return 2 team records", "Returned merged CSV metrics", "PASS"],
        ["TC-05", "Vite Build", "npm run build", "Production bundle compilation", "Built in 1.17s with 0 errors", "PASS"]
    ]
    t_tc = doc.add_table(rows=1, cols=6)
    style_table(t_tc, tc_widths, tc_headers, tc_data)

    # -------------------------------------------------------------
    # 25. KPIs
    # -------------------------------------------------------------
    add_h1("25. Key Performance Indicators (KPIs)")
    add_bullet("• Technical KPIs: ", "5-Fold CV Accuracy (74.38%), API Response Time (<30ms), Frontend Build Time (1.17s), JS Bundle Size (212 kB).")
    add_bullet("• Academic KPIs: ", "Historical Backtesting Score (72.7%), Data Completeness (100% across 12 World Cups).")

    # -------------------------------------------------------------
    # 26. SUCCESS METRICS
    # -------------------------------------------------------------
    add_h1("26. Success Metrics")
    add_bullet("1. Zero Runtime Errors: ", "Clean console logs and error-free execution across all 7 frontend page views.")
    add_bullet("2. State Persistence: ", "100% reliable LocalStorage client prediction saving for Raj's picks.")
    add_bullet("3. Accurate Data Rendering: ", "Factual alignment with official ICC World Cup archives.")

    # -------------------------------------------------------------
    # 27. ROADMAP
    # -------------------------------------------------------------
    add_h1("27. Development Roadmap")
    rm_headers = ["Phase", "Milestones & Scope", "Target Deliverables", "Status"]
    rm_widths = [0.8, 3.2, 1.8, 0.7]
    rm_data = [
        ["MVP", "1983-2027 history dataset, 10-feature ML model, Raj vs AI engine, FastAPI, React dark UI.", "Fully functional local web app", "COMPLETED"],
        ["V1", "Improved feature weights, expanded team stats, exportable PDF report.", "Enhanced analytical dashboard", "Current"],
        ["V2", "Live ICC webhook API integration, ball-by-ball win probability graph, relational Postgres DB.", "Real-time sports data portal", "Planned"],
        ["V3", "10,000 Monte Carlo match simulations, mobile React Native app, multi-user prediction challenges.", "Enterprise sports AI platform", "Planned"]
    ]
    t_rm = doc.add_table(rows=1, cols=4)
    style_table(t_rm, rm_widths, rm_headers, rm_data)

    # -------------------------------------------------------------
    # 28. RISKS & MITIGATION
    # -------------------------------------------------------------
    add_h1("28. Risk & Mitigation Matrix")
    rk_headers = ["Risk Description", "Impact", "Probability", "Mitigation Strategy"]
    rk_widths = [2.0, 0.8, 0.8, 2.9]
    rk_data = [
        ["Small Historical Sample Size (11 WCs)", "High", "High", "Generated synthetic jitter training candidates to train robust classifiers."],
        ["Knockout Match Non-Determinism", "Medium", "High", "Calculated Softmax probability distributions rather than hard binary outputs."],
        ["Backend Service Offline", "High", "Low", "Built offline fallback dataset in frontend api.js to ensure 100% UI uptime."]
    ]
    t_rk = doc.add_table(rows=1, cols=4)
    style_table(t_rk, rk_widths, rk_headers, rk_data)

    # -------------------------------------------------------------
    # 29. FUTURE ENHANCEMENTS
    # -------------------------------------------------------------
    add_h1("29. Future Enhancements")
    add_bullet("1. Real-Time Webhook Integration: ", "Connecting live match score APIs to update win probabilities ball-by-ball.")
    add_bullet("2. Player-Level Elo Ratings: ", "Incorporating individual batting/bowling Elo ratings and squad injury status.")
    add_bullet("3. Monte Carlo Tournament Simulator: ", "Simulating 10,000 automated tournament match iterations.")

    # -------------------------------------------------------------
    # 30. APPENDIX
    # -------------------------------------------------------------
    add_h1("30. Appendix")
    
    add_h2("A. Technology Stack Reference")
    p_app_a = doc.add_paragraph()
    p_app_a.add_run("React 18.2, Vite 5.1, Tailwind CSS 3.4, Python 3.10+, FastAPI 0.141, Scikit-Learn 1.6, Pandas 2.2, NumPy 2.1.")

    add_h2("B. Complete Project Directory Tree")
    p_app_b = doc.add_paragraph()
    r_tree = p_app_b.add_run(
        "cricket-analytics-2027/\n"
        "├── data/ (world_cup_history.json, team_wc_performance.csv, recent_odi_stats.csv)\n"
        "├── ml/ (feature_engineering.py, train_model.py)\n"
        "├── models/ (best_model.pkl, scaler.pkl, prediction_2027.json)\n"
        "├── backend/ (main.py, requirements.txt)\n"
        "├── frontend/ (index.html, vite.config.js, package.json, src/App.jsx, src/components/, src/pages/)\n"
        "└── README.md"
    )
    r_tree.font.name = 'Consolas'
    r_tree.font.size = Pt(8.5)

    add_h2("C. Dataset Schema Breakdown")
    p_app_c = doc.add_paragraph()
    p_app_c.add_run("team_wc_performance.csv: Team, Titles, Finals_Appearances, Semi_Finals, Matches_Played, Matches_Won, Matches_Lost, Win_Percentage, Batting_Avg, Bowling_Avg, Net_Run_Rate, Knockout_Win_Pct.")

    add_h2("D. Model Features")
    p_app_d = doc.add_paragraph()
    p_app_d.add_run("10 Features: feat_wc_win_pct, feat_title_exp, feat_knockout_clutch, feat_recent_win_pct, feat_batting_strength, feat_bowling_strength, feat_consistency, feat_host_adaptability, feat_h2h_top5, feat_power_index.")

    add_h2("E. API Reference")
    p_app_e = doc.add_paragraph()
    p_app_e.add_run("GET / | GET /api/health | GET /api/prediction/2027 | GET /api/history | GET /api/teams | GET /api/teams/compare | GET /api/model/metrics.")

    add_h2("F. UI Screenshots Placeholders")
    p_app_f = doc.add_paragraph()
    p_app_f.add_run("[Placeholder: Solid Black Cricket Splash Screen (2.5s)]\n[Placeholder: Home Dashboard & 2027 Winner Card]\n[Placeholder: Raj vs AI Prediction Matrix & Scoreboard]\n[Placeholder: Team Side-by-Side Comparison Tool]\n[Placeholder: AI Analysis & Feature Importance Bar Chart]")

    add_h2("G. Technical Glossary")
    p_app_g = doc.add_paragraph()
    p_app_g.add_run("• ODI: One Day International cricket match (50 overs per side).\n• Softmax: Mathematical function that normalizes a vector of K real numbers into a probability distribution summing to 1.0.\n• Stratified K-Fold CV: Cross-validation technique that preserves target class proportions across all folds.\n• FastAPI: Modern, high-performance web framework for building APIs with Python.")

    # Save to both target locations
    path1 = "/Users/rajmalviya/Desktop/cricket-analytics-2027/PROJECT_DOCUMENTATION.docx"
    path2 = "/Users/rajmalviya/Desktop/cricket-analytics-2027/ODI_World_Cup_2027_Academic_Documentation.docx"
    doc.save(path1)
    doc.save(path2)
    print(f"Saved documentation to:\n  - {path1}\n  - {path2}")

if __name__ == "__main__":
    create_document()
