import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_document():
    doc = Document()
    
    # Page setup - Margins (1 inch all sides)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Color Palette
    PRIMARY_COLOR = RGBColor(16, 185, 129)    # Emerald Green #10B981
    SECONDARY_COLOR = RGBColor(245, 158, 11)  # Trophy Gold #F59E0B
    DARK_BG = RGBColor(15, 23, 42)            # Slate Dark #0F172A
    TEXT_COLOR = RGBColor(30, 41, 59)          # Slate Text #1E293B
    MUTED_COLOR = RGBColor(100, 116, 139)      # Slate Muted #64748B

    # Helper function for setting cell shading
    def set_cell_background(cell, fill_hex):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # Helper function for setting cell margins/padding
    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
            node = OxmlElement(margin)
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    # Set Default Font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = TEXT_COLOR

    # Function to add styled headings
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = SECONDARY_COLOR
        return p

    def add_heading_3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = DARK_BG
        return p

    def add_bullet(p_or_text, bold_prefix="", text=""):
        if isinstance(p_or_text, str):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(3)
            if bold_prefix:
                r_bold = p.add_run(bold_prefix)
                r_bold.bold = True
            p.add_run(text)
        else:
            p = p_or_text

    # Function to format tables nicely
    def style_table(table, col_widths, headers, data):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Header row
        hdr_cells = table.rows[0].cells
        for i, header_text in enumerate(headers):
            hdr_cells[i].text = header_text
            set_cell_background(hdr_cells[i], "0F172A") # Dark Navy Header
            set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10.5)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Data rows
        for row_idx, row_data in enumerate(data):
            row_cells = table.add_row().cells
            bg_color = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
            for i, val in enumerate(row_data):
                row_cells[i].text = str(val)
                set_cell_background(row_cells[i], bg_color)
                set_cell_margins(row_cells[i], top=100, bottom=100, left=150, right=150)
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(10)
                        run.font.color.rgb = TEXT_COLOR
                        
        # Column widths
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    # -------------------------------------------------------------
    # DOCUMENT COVER / TITLE BLOCK
    # -------------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(20)
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run("ODI World Cup 2027 AI Analytics & Prediction System")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = PRIMARY_COLOR

    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_before = Pt(0)
    subtitle_p.paragraph_format.space_after = Pt(12)
    sub_run = subtitle_p.add_run("Raj vs AI — ODI World Cup Prediction Platform")
    sub_run.font.name = 'Calibri'
    sub_run.font.size = Pt(16)
    sub_run.font.bold = True
    sub_run.font.color.rgb = SECONDARY_COLOR

    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_before = Pt(0)
    meta_p.paragraph_format.space_after = Pt(24)
    meta_run = meta_p.add_run("Author: Raj Malviya  |  Technical Documentation & B.Tech Project Report")
    meta_run.font.name = 'Calibri'
    meta_run.font.size = Pt(11)
    meta_run.font.italic = True
    meta_run.font.color.rgb = MUTED_COLOR

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # -------------------------------------------------------------
    # 1. PROJECT OVERVIEW
    # -------------------------------------------------------------
    add_heading_1("1. Project Overview")
    
    p = doc.add_paragraph()
    p.add_run("The ")
    r = p.add_run("ODI World Cup 2027 AI Analytics & Prediction System")
    r.bold = True
    p.add_run(" is an advanced full-stack data analytics platform and machine learning system engineered to analyze 40 years of ICC Men's ODI World Cup historical records (1983–2023), team strength indicators, and regional venue adaptability metrics. The primary goal of the system is to estimate win probabilities for the upcoming 2027 ODI World Cup (hosted across South Africa, Zimbabwe, and Namibia) while providing an interactive head-to-head gamified experience called ")
    r2 = p.add_run("Raj vs AI")
    r2.bold = True
    p.add_run(".")

    p2 = doc.add_paragraph()
    p2.add_run("In traditional sports journalism, tournament predictions are frequently based on subjective commentary or biased fan opinions. This project solves that problem by introducing a quantitative, data-driven approach. By extracting 10 empirical feature parameters from 11 past World Cup editions (556 total matches) and recent international ODI performance data (2024–2026), the system trains and evaluates supervised classification models to compute Softmax-normalized probability distributions across all major participating teams.")

    p3 = doc.add_paragraph()
    p3.add_run("A central feature of the platform is the ")
    p3.add_run("Raj vs AI").bold = True
    p3.add_run(" prediction engine. It allows the user (Raj) to select manual predictions for every World Cup tournament from 1983 to 2027. The system dynamically evaluates both Raj's selections and the AI's machine learning predictions against actual historical champions (1983–2023), maintaining real-time accuracy scoreboards, leader standings, and year-by-year probability breakdowns.")

    # -------------------------------------------------------------
    # 2. PROBLEM STATEMENT
    # -------------------------------------------------------------
    add_heading_1("2. Problem Statement")
    
    p = doc.add_paragraph()
    p.add_run("Predicting the champion of the ODI World Cup is a multifaceted challenge. Tournament outcomes are governed by complex interconnected factors, including:")
    
    add_bullet("1. Historical World Cup Pedigree: ", "Past titles, finals appearances, and long-term tournament experience.", "")
    add_bullet("2. Knockout Stage Clutch Conversion: ", "The ability of a team to perform under intense pressure in semi-finals and finals.", "")
    add_bullet("3. Team Composition Balance: ", "The dual strength of batting depth (runs per wicket, run rate) and bowling efficiency (economy rate, wickets per match).", "")
    add_bullet("4. Recent Bilateral Form: ", "International ODI win percentage over the preceding 3-year cycle (2024–2026).", "")
    add_bullet("5. Host Venue Adaptability: ", "Condition compatibility in African sub-continent pitches (South Africa, Zimbabwe, Namibia host venues for 2027).", "")

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(6)
    p2.add_run("This project addresses the problem by designing a structured machine learning pipeline that standardizes historical and recent statistics into a 10-dimensional feature vector, trains supervised classifiers, and calculates non-deterministic Softmax win probabilities while allowing direct comparison against human intuition (Raj's predictions).")

    # -------------------------------------------------------------
    # 3. PROJECT OBJECTIVES
    # -------------------------------------------------------------
    add_heading_1("3. Project Objectives")
    
    p = doc.add_paragraph()
    p.add_run("The key technical and functional objectives of the project are:")

    add_bullet("• Data Collection & Structuring: ", "Compile and standardize historical ICC ODI World Cup data from 1983 to 2023, encompassing 11 tournament editions, match scorelines, top scorers, leading wicket-takers, and cumulative team statistics.")
    add_bullet("• Feature Engineering: ", "Develop a 10-feature quantitative dataset capturing WC win rates, title experience, knockout clutch ratings, recent 3-year win rates, batting ratings, bowling ratings, consistency scores, and African host venue adaptability.")
    add_bullet("• Supervised ML Modeling: ", "Train, validate, and compare 4 classification algorithms (Random Forest, Gradient Boosting, Extra Trees, and Logistic Regression) using 5-fold Stratified Cross-Validation.")
    add_bullet("• Probability Calibration: ", "Generate Softmax-normalized win probability scores for all 10 major participating teams for the 2027 World Cup.")
    add_bullet("• Raj vs AI Prediction Matrix: ", "Build an interactive dual-prediction engine covering every World Cup from 1983 to 2027, with client-side LocalStorage state persistence for Raj's picks.")
    add_bullet("• Historical Backtesting: ", "Evaluate the AI model across all past World Cups (1983–2023) to compute an empirical historical accuracy score (achieving 72.7% backtest accuracy).")
    add_bullet("• RESTful API Architecture: ", "Expose high-performance FastAPI backend endpoints to deliver dataset records, model metrics, team comparisons, and predictions.")
    add_bullet("• Clean Responsive Frontend: ", "Design a sleek B.Tech project UI using React 18, Vite, and Tailwind CSS featuring a solid black cricket splash screen and responsive analytical views.")

    # -------------------------------------------------------------
    # 4. KEY FEATURES
    # -------------------------------------------------------------
    add_heading_1("4. Key Features")

    features = [
        ("Solid Black Full-Screen Cricket Splash Screen", "Presents a 100% solid black full-screen overlay (fixed inset-0 w-screen h-screen z-[999999] bg-black) on initial page load/refresh. Shows a bowling ball animation toward wickets, bat impact, and a 4-step text sequence ('ODI WORLD CUP 2027' -> 'AI ANALYTICS' -> 'RAJ VS AI' -> Progress bar) for ~2.5 seconds before cleanly revealing the website.", "Implemented in CricketSplashScreen.jsx. Rendered exclusively when showSplash is true in App.jsx, ensuring 0% of the underlying app is visible during loading."),
        ("Home Dashboard View", "Provides a comprehensive project overview featuring the 2027 AI Predicted Champion card (Australia 31.3%), Top 5 contenders leaderboard, key tournament counters (11 World Cups, 556 matches, 6 unique champions), and a 1983-2023 champions honor roll.", "Implemented in Home.jsx. Fetches live prediction data from the FastAPI backend or local fallback services."),
        ("World Cup History Archive", "Displays a complete table of all 12 World Cup editions (1983–2027) with host countries, actual winners, runner-ups, total matches, and interactive drill-down inspection modals detailing final scorelines, top scorers, and summaries.", "Implemented in History.jsx. Supports real-time text searching across years, teams, and hosts."),
        ("Raj vs AI Complete Prediction Engine", "The centerpiece feature of the application. Allows Raj to select predictions for every World Cup from 1983 to 2027, maintains persistent state in LocalStorage, calculates live scoreboards, tracks leader standings, and displays 2027 pick comparisons.", "Implemented in RajVsAi.jsx. Automatically compares Raj's pick and AI's pick against actual historical champions (1983-2023) and marks 2027 as TBD/Pending."),
        ("Team Analysis & Comparison Tool", "Allows users to select between 2 and 4 teams (e.g., India, Australia, South Africa) and compare them side-by-side across 9 quantitative metrics including titles, win %, knockout clutch, batting rating, bowling rating, and African host rating.", "Implemented in TeamAnalytics.jsx. Renders a clean grid layout for side-by-side statistical inspection."),
        ("AI Analysis & Model Metrics View", "Exposes the machine learning pipeline performance, cross-validation metrics (Accuracy, Precision, Recall, F1-Score), algorithm comparison table, 10 feature importance weights, and full 10-team probability rankings.", "Implemented in AiAnalysis.jsx. Reads cross-validated model outputs from prediction_2027.json."),
        ("2027 Prediction Special Highlight", "Dedicated 2027 World Cup view displaying the highlighted Raj's Pick vs AI's Pick comparison card, agreement status ('Raj and AI agree/disagree'), and the prominent 2027 team probability ranking table.", "Implemented in Prediction2027.jsx."),
        ("Data Sources & Disclaimer View", "Documents data references (ICC Archives, ESPNcricinfo, ODI Rankings) and presents a concise statistical machine learning prediction disclaimer.", "Implemented in DataSources.jsx. Fully replaces all speculative or philosophical AI text.")
    ]

    for fname, fdesc, fimpl in features:
        add_heading_2(fname)
        p = doc.add_paragraph()
        p.add_run("What it does: ").bold = True
        p.add_run(fdesc + "\n")
        p.add_run("How it works & Implementation: ").bold = True
        p.add_run(fimpl)

    # -------------------------------------------------------------
    # 5. RAJ VS AI FEATURE
    # -------------------------------------------------------------
    add_heading_1("5. Raj vs AI Feature")

    p = doc.add_paragraph()
    p.add_run("The ")
    p.add_run("Raj vs AI").bold = True
    p.add_run(" feature provides an interactive, comparative prediction matrix for every ODI World Cup edition from 1983 to 2027. The system workflow follows a clear sequential logic:")

    p_flow = doc.add_paragraph()
    p_flow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_flow = p_flow.add_run("Raj Selects Prediction  →  AI Generates Prediction  →  Raj vs AI Comparison  →  Historical Actual Winner  →  Correct / Incorrect Result")
    r_flow.bold = True
    r_flow.font.color.rgb = PRIMARY_COLOR

    add_heading_2("Workflow Breakdown")
    add_bullet("1. Prediction Selection: ", "The user selects a tournament year (1983–2027) and picks a predicted winner from a dropdown menu. Clicking 'Submit Prediction' saves the choice in client-side LocalStorage (raj_wc_predictions).", "")
    add_bullet("2. AI Prediction Generation: ", "For 1983–2023, the system retrieves the backtested ML model prediction for that tournament year. For 2027, it retrieves the model's top predicted team (Australia, 31.3%).", "")
    add_bullet("3. Comparison & Verification: ", "For historical years (1983–2023), Raj's pick and AI's pick are compared against the Actual Winner. If Raj's Pick == Actual Winner, Raj gets 'Correct'; otherwise 'Incorrect'. Similarly for the AI.", "")
    add_bullet("4. 2027 Special Handling: ", "For the upcoming 2027 tournament, the Actual Winner is designated as 'TBD'. Consequently, Raj Result and AI Result are displayed as 'Pending', while a special banner indicates whether Raj and AI agree or disagree.", "")
    add_bullet("5. Scoreboard & Leaderboard: ", "The page computes total Correct, Incorrect, and Accuracy % for both Raj and AI across 1983–2023 (excluding 2027). A dynamic header banner announces: 'AI is currently leading!', 'Raj is currently leading!', or 'It's a tie'.", "")

    # -------------------------------------------------------------
    # 6. ODI WORLD CUP HISTORICAL DATA
    # -------------------------------------------------------------
    add_heading_1("6. ODI World Cup Historical Data")

    p = doc.add_paragraph()
    p.add_run("The project dataset includes complete historical records for all 12 World Cup editions from 1983 to 2027. Below is the authoritative historical tournament log stored in ")
    p.add_run("data/world_cup_history.json").bold = True
    p.add_run(":")

    wc_headers = ["Year", "Host Country", "Actual Winner", "Runner-Up", "Matches", "Top Scorer", "Top Wicket Taker"]
    wc_col_widths = [0.8, 1.3, 1.1, 1.1, 0.7, 1.4, 1.4]
    wc_data = [
        ["1983", "England", "India", "West Indies", "27", "David Gower (ENG, 384)", "Roger Binny (IND, 18)"],
        ["1987", "India & Pakistan", "Australia", "England", "27", "Graham Gooch (ENG, 471)", "Craig McDermott (AUS, 18)"],
        ["1992", "Australia & NZ", "Pakistan", "England", "39", "Martin Crowe (NZ, 456)", "Wasim Akram (PAK, 18)"],
        ["1996", "IND, PAK & SL", "Sri Lanka", "Australia", "37", "Sachin Tendulkar (IND, 523)", "Anil Kumble (IND, 15)"],
        ["1999", "England", "Australia", "Pakistan", "42", "Rahul Dravid (IND, 461)", "S. Warne / G. Allott (20)"],
        ["2003", "South Africa", "Australia", "India", "54", "Sachin Tendulkar (IND, 673)", "Chaminda Vaas (SL, 23)"],
        ["2007", "West Indies", "Australia", "Sri Lanka", "51", "Matthew Hayden (AUS, 659)", "Glenn McGrath (AUS, 26)"],
        ["2011", "IND, SL & BAN", "India", "Sri Lanka", "49", "T. Dilshan (SL, 500)", "Z. Khan / S. Afridi (21)"],
        ["2015", "Australia & NZ", "Australia", "New Zealand", "49", "Martin Guptill (NZ, 547)", "M. Starc / T. Boult (22)"],
        ["2019", "England & Wales", "England", "New Zealand", "48", "Rohit Sharma (IND, 648)", "Mitchell Starc (AUS, 27)"],
        ["2023", "India", "Australia", "India", "48", "Virat Kohli (IND, 765)", "Mohammed Shami (IND, 24)"],
        ["2027", "SA, ZIM & NAM", "TBD", "TBD", "54", "TBD", "TBD"]
    ]

    t_wc = doc.add_table(rows=1, cols=7)
    style_table(t_wc, wc_col_widths, wc_headers, wc_data)

    # -------------------------------------------------------------
    # 7. DATASET
    # -------------------------------------------------------------
    add_heading_1("7. Dataset Specification")

    p = doc.add_paragraph()
    p.add_run("The project relies on structured JSON and CSV datasets located in the ")
    p.add_run("data/").bold = True
    p.add_run(" and ")
    p.add_run("models/").bold = True
    p.add_run(" directories:")

    ds_headers = ["File Name", "Location", "File Type", "Purpose & Contents"]
    ds_widths = [1.8, 1.0, 1.0, 3.2]
    ds_data = [
        ["world_cup_history.json", "data/", "JSON", "Stores tournament-level records for 1983-2027 including year, host, winner, runner-up, scorelines, top scorers, top wicket-takers, summaries, and AI backtest predictions."],
        ["team_wc_performance.csv", "data/", "CSV", "Cumulative World Cup team stats (1983-2023). Columns: Team, Titles, Finals_Appearances, Semi_Finals, Matches_Played, Matches_Won, Matches_Lost, Win_Percentage, Batting_Avg, Bowling_Avg, Net_Run_Rate, Knockout_Win_Pct."],
        ["recent_odi_stats.csv", "data/", "CSV", "Recent 3-year ODI form (2024-2026) and 2027 venue factors. Columns: Team, ICC_Rank, ICC_Rating, Recent_Win_Pct, Batting_Rating, Bowling_Rating, Consistency_Score, African_Host_Rating, H2H_Top5_Win_Pct."],
        ["prediction_2027.json", "models/", "JSON", "Exported Machine Learning outputs including model_used, CV metrics, top_5_teams, all_predictions, feature_importances, explanation, and historical_backtest array."]
    ]

    t_ds = doc.add_table(rows=1, cols=4)
    style_table(t_ds, ds_widths, ds_headers, ds_data)

    # -------------------------------------------------------------
    # 8. MACHINE LEARNING
    # -------------------------------------------------------------
    add_heading_1("8. Machine Learning Model & Evaluation")

    p = doc.add_paragraph()
    p.add_run("The machine learning module (")
    p.add_run("ml/train_model.py").bold = True
    p.add_run(") trains and evaluates 4 supervised classification algorithms using 5-fold Stratified Cross-Validation on a synthetic historical dataset generated from empirical feature vectors.")

    add_heading_2("Algorithm Performance Comparison")
    
    ml_headers = ["Algorithm", "Validation Accuracy", "Precision", "Recall", "F1-Score", "Selection Status"]
    ml_widths = [1.8, 1.2, 1.0, 1.0, 1.0, 1.0]
    ml_data = [
        ["Logistic Regression", "74.38%", "66.46%", "46.27%", "0.5070", "SELECTED BEST"],
        ["Gradient Boosting", "71.88%", "54.72%", "46.66%", "0.4982", "Evaluated"],
        ["Random Forest", "70.63%", "52.97%", "44.22%", "0.4552", "Evaluated"],
        ["Extra Trees", "71.88%", "62.86%", "35.03%", "0.3992", "Evaluated"]
    ]

    t_ml = doc.add_table(rows=1, cols=6)
    style_table(t_ml, ml_widths, ml_headers, ml_data)

    p_reason = doc.add_paragraph()
    p_reason.paragraph_format.space_before = Pt(8)
    p_reason.add_run("Why Logistic Regression was Selected: ").bold = True
    p_reason.add_run("Logistic Regression achieved the highest cross-validated F1-score (0.5070) and validation accuracy (74.38%). Its linear log-odds decision boundary provides smooth, well-calibrated probability outputs when transformed via Softmax, preventing over-confident extreme predictions.")

    add_heading_2("10 Engineered Input Features & Weights")
    
    feat_headers = ["Feature Code", "Feature Label", "Importance Weight", "Description"]
    feat_widths = [1.8, 2.2, 1.0, 2.0]
    feat_data = [
        ["feat_knockout_clutch", "Knockout Stage Win Rate", "23.83%", "Historical win % in WC semi-finals and finals."],
        ["feat_h2h_top5", "Head-to-Head vs Top 5", "17.00%", "Win % against top 5 ICC ranked ODI nations."],
        ["feat_consistency", "Tournament Consistency", "13.65%", "Consistency rating across multi-game tournaments."],
        ["feat_host_adaptability", "African Host Advantage", "10.77%", "Rating in SA, ZIM, NAM venue conditions."],
        ["feat_bowling_strength", "Bowling Efficiency", "10.05%", "Composite bowling economy and wicket-taking rate."],
        ["feat_wc_win_pct", "Historical WC Win Rate", "8.37%", "Overall career win % in all World Cup matches."],
        ["feat_recent_win_pct", "Recent 3-Year ODI Form", "5.91%", "Win % in bilateral ODIs during 2024–2026."],
        ["feat_batting_strength", "Batting Depth & Average", "5.90%", "Composite batting average and run-rate rating."],
        ["feat_title_exp", "Titles & Finals Record", "3.58%", "Scaled weight of titles and finals appearances."],
        ["feat_power_index", "Composite Power Index", "0.94%", "Weighted combination of all core sub-indices."]
    ]

    t_feat = doc.add_table(rows=1, cols=4)
    style_table(t_feat, feat_widths, feat_headers, feat_data)

    # -------------------------------------------------------------
    # 9. ML PIPELINE
    # -------------------------------------------------------------
    add_heading_1("9. Machine Learning Pipeline")

    p = doc.add_paragraph()
    p.add_run("The complete data and machine learning workflow executes through 7 structured stages:")

    add_bullet("Stage 1: Raw Data Loading — ", "Reads team_wc_performance.csv and recent_odi_stats.csv from data/.", "")
    add_bullet("Stage 2: Feature Engineering — ", "Computes 10 normalized features (0.0 to 1.0 scale) in ml/feature_engineering.py.", "")
    add_bullet("Stage 3: Synthetic Historical Dataset Creation — ", "Generates historical tournament candidates with realistic variance for training.", "")
    add_bullet("Stage 4: Standard Scaling & Cross-Validation — ", "Applies StandardScaler and 5-fold Stratified K-Fold CV across 4 algorithms.", "")
    add_bullet("Stage 5: Model Selection & Refitting — ", "Selects Logistic Regression based on F1-score and refits on the complete dataset.", "")
    add_bullet("Stage 6: Softmax Temperature Probability Calculation — ", "Transforms decision function log-odds into normalized team win probabilities.", "")
    add_bullet("Stage 7: Artifact Export — ", "Saves best_model.pkl, scaler.pkl, and prediction_2027.json for FastAPI backend consumption.", "")

    # -------------------------------------------------------------
    # 10. BACKEND API
    # -------------------------------------------------------------
    add_heading_1("10. Backend Architecture & API Specification")

    p = doc.add_paragraph()
    p.add_run("The backend is built with ")
    p.add_run("Python FastAPI").bold = True
    p.add_run(" and hosted via Uvicorn (")
    p.add_run("backend/main.py").bold = True
    p.add_run("). It exposes REST API endpoints with CORS middleware enabled for frontend communication.")

    api_headers = ["Endpoint", "HTTP Method", "Query Parameters", "Purpose & Return Data"]
    api_widths = [1.8, 1.0, 1.4, 2.8]
    api_data = [
        ["/", "GET", "None", "Returns API status, application name, and documentation link."],
        ["/api/health", "GET", "None", "Returns system health check status ('healthy')."],
        ["/api/prediction/2027", "GET", "None", "Returns 2027 ML predictions, Softmax team probabilities, feature importances, explanation narrative, and 1983-2027 backtest array."],
        ["/api/history", "GET", "year (optional int)", "Returns historical World Cup records from 1983 to 2027. Supports filtering by year."],
        ["/api/teams", "GET", "None", "Returns merged historical World Cup performance and recent ODI statistics for all teams."],
        ["/api/teams/compare", "GET", "teams (str, default: 'India,Australia')", "Compares two or more teams side-by-side on historical and recent metrics."],
        ["/api/model/metrics", "GET", "None", "Returns validation metrics across all 4 trained ML algorithms and feature weights."]
    ]

    t_api = doc.add_table(rows=1, cols=4)
    style_table(t_api, api_widths, api_headers, api_data)

    # -------------------------------------------------------------
    # 11. FRONTEND ARCHITECTURE
    # -------------------------------------------------------------
    add_heading_1("11. Frontend Architecture")

    p = doc.add_paragraph()
    p.add_run("The frontend application is built using ")
    p.add_run("React 18").bold = True
    p.add_run(", ")
    p.add_run("Vite 5").bold = True
    p.add_run(", and ")
    p.add_run("Tailwind CSS 3").bold = True
    p.add_run(". It follows a modular component structure located in ")
    p.add_run("frontend/src/").bold = True
    p.add_run(":")

    add_bullet("• App.jsx: ", "Main entry container managing state for showSplash, activeTab, predictionData, historyData, and dataLoading. Guarantees that only CricketSplashScreen renders during loading.")
    add_bullet("• CricketSplashScreen.jsx: ", "Solid black overlay (#000000) rendering bowling SVG animations, wickets, bat hit impact, and a 4-step sequential text reveal.")
    add_bullet("• Navbar.jsx & Footer.jsx: ", "Navigation header with 7 active tabs and footer displaying 'Built by Raj Malviya'.")
    add_bullet("• RajVsAi.jsx: ", "Interactive Raj vs AI prediction engine with year selection, LocalStorage state persistence, scoreboard, leader banner, and 2027 pick comparison.")
    add_bullet("• Home.jsx, History.jsx, TeamAnalytics.jsx, AiAnalysis.jsx, Prediction2027.jsx, DataSources.jsx: ", "Dedicated analytical page views.")
    add_bullet("• api.js: ", "API client handling HTTP requests to http://localhost:8000/api with built-in fallback datasets.")

    # -------------------------------------------------------------
    # 12. TECHNOLOGY STACK
    # -------------------------------------------------------------
    add_heading_1("12. Technology Stack")

    tech_headers = ["Technology", "Category", "Version", "Role in Project"]
    tech_widths = [1.5, 1.2, 1.0, 3.3]
    tech_data = [
        ["React", "Frontend Framework", "18.2.0", "Component-based user interface rendering and state management."],
        ["Vite", "Frontend Build Tool", "5.1.6", "Ultra-fast development server and production bundler (1.17s build)."],
        ["Tailwind CSS", "Styling Framework", "3.4.1", "Utility-first CSS styling for clean, dark theme B.Tech project aesthetic."],
        ["Lucide React", "Icon System", "0.344.0", "Clean vector icons for navigation, trophies, and status badges."],
        ["Python", "Backend & ML Runtime", "3.10+", "Execution environment for machine learning training and API serving."],
        ["FastAPI", "Backend Web Framework", "0.141.1", "High-performance asynchronous REST API framework."],
        ["Uvicorn", "ASGI Web Server", "0.52.4", "Production ASGI server hosting backend endpoints on port 8000."],
        ["Scikit-Learn", "Machine Learning Library", "1.6.1", "Model training, StandardScaler, KFold cross-validation, and metrics."],
        ["Pandas", "Data Manipulation", "2.2.3", "Dataframe loading, preprocessing, merging CSV/JSON datasets."],
        ["NumPy", "Numerical Computing", "2.1.3", "Array manipulations, jitter matrix generation, and Softmax scaling."],
        ["LocalStorage API", "Web Storage API", "Browser Native", "Client-side persistence of Raj's manual predictions across all 12 years."]
    ]

    t_tech = doc.add_table(rows=1, cols=4)
    style_table(t_tech, tech_widths, tech_headers, tech_data)

    # -------------------------------------------------------------
    # 13. PROJECT FOLDER STRUCTURE
    # -------------------------------------------------------------
    add_heading_1("13. Project Folder Structure")

    p_struct = doc.add_paragraph()
    p_struct.add_run("The repository is organized into a clean, modular folder hierarchy:")

    struct_text = (
        "cricket-analytics-2027/\
"
        "├── data/\
"
        "│   ├── world_cup_history.json      # 1983-2027 tournament records & AI backtest data\
"
        "│   ├── team_wc_performance.csv     # Historical team WC stats (1983-2023)\
"
        "│   └── recent_odi_stats.csv        # Recent 3-year form & African venue ratings\
"
        "├── ml/\
"
        "│   ├── feature_engineering.py      # Feature extraction & 10-vector normalization\
"
        "│   └── train_model.py              # ML cross-validation, backtesting & export\
"
        "├── models/\
"
        "│   ├── best_model.pkl              # Saved Logistic Regression model binary\
"
        "│   ├── scaler.pkl                  # Saved StandardScaler object\
"
        "│   └── prediction_2027.json        # Exported JSON predictions & metrics\
"
        "├── backend/\
"
        "│   ├── main.py                     # FastAPI server with 7 REST API endpoints\
"
        "│   └── requirements.txt            # Python dependencies\
"
        "├── frontend/\
"
        "│   ├── index.html                  # HTML5 entrypoint\
"
        "│   ├── vite.config.js              # Vite server & proxy configuration\
"
        "│   ├── package.json                # React dependencies\
"
        "│   └── src/\
"
        "│       ├── App.jsx                 # Main state container & loader router\
"
        "│       ├── index.css               # Tailwind directives & dark theme styling\
"
        "│       ├── components/\
"
        "│       │   ├── CricketSplashScreen.jsx  # Solid black overlay loader\
"
        "│       │   ├── Navbar.jsx               # 7-tab navigation header\
"
        "│       │   └── Footer.jsx               # Footer with 'Built by Raj Malviya'\
"
        "│       ├── pages/\
"
        "│       │   ├── Home.jsx                 # Overview dashboard\
"
        "│       │   ├── History.jsx              # 1983-2027 history & inspection\
"
        "│       │   ├── RajVsAi.jsx              # Dual prediction engine & scoreboard\
"
        "│       │   ├── TeamAnalytics.jsx        # Side-by-side team comparison\
"
        "│       │   ├── AiAnalysis.jsx           # ML metrics & feature weights\
"
        "│       │   ├── Prediction2027.jsx       # 2027 Raj vs AI pick special view\
"
        "│       │   └── DataSources.jsx          # References & ML disclaimer\
"
        "│       └── services/\
"
        "│           └── api.js                   # API client with fallback data\
"
        "└── README.md                       # Documentation & run instructions"
    )

    p_code = doc.add_paragraph()
    p_code.paragraph_format.space_before = Pt(4)
    p_code.paragraph_format.space_after = Pt(8)
    r_code = p_code.add_run(struct_text)
    r_code.font.name = 'Consolas'
    r_code.font.size = Pt(8.5)
    r_code.font.color.rgb = DARK_BG

    # -------------------------------------------------------------
    # 14. WEBSITE WORKFLOW
    # -------------------------------------------------------------
    add_heading_1("14. Website Workflow")

    p = doc.add_paragraph()
    p.add_run("The user interaction journey through the website follows a clean 7-step path:")

    add_bullet("1. Page Launch: ", "User opens http://localhost:3000. App renders ONLY the solid black CricketSplashScreen.", "")
    add_bullet("2. Splash Animation (2.5s): ", "Cricket ball animation plays -> Text sequence reveals 'ODI WORLD CUP 2027' -> 'AI ANALYTICS' -> 'RAJ VS AI' -> Fades out.", "")
    add_bullet("3. Home Dashboard: ", "Main site opens. User views 2027 predicted winner (Australia 31.3%), top 5 contenders, and quick historical stats.", "")
    add_bullet("4. Raj vs AI Predictions: ", "User navigates to 'Raj vs AI' tab. Selects a World Cup year (1983–2027) and picks a predicted winner.", "")
    add_bullet("5. Scoreboard & Leader Update: ", "System saves prediction in LocalStorage, updates Raj's score, and displays 'AI is currently leading' or 'Raj is currently leading'.", "")
    add_bullet("6. Team & AI Analysis: ", "User compares team metrics or inspects ML feature importance weights and cross-validation accuracy.", "")
    add_bullet("7. Data Sources Inspection: ", "User reviews references and prediction disclaimer.")

    # -------------------------------------------------------------
    # 15. 2027 PREDICTION
    # -------------------------------------------------------------
    add_heading_1("15. 2027 World Cup Prediction Details")

    p = doc.add_paragraph()
    p.add_run("For the 2027 ODI World Cup (hosted jointly by South Africa, Zimbabwe, and Namibia), the AI model computes the following Softmax probability ranking:")

    prob_headers = ["Rank", "Team Name", "Win Probability %", "Key Contributing Factors"]
    prob_widths = [0.8, 1.5, 1.5, 3.2]
    prob_data = [
        ["#1", "Australia", "31.3%", "75% Knockout Win Rate, 6 WC Titles, 82.0 African Host Rating."],
        ["#2", "India", "23.9%", "72.5% Recent Win Rate (2024-2026), 88.5 Batting Rating, 2 WC Titles."],
        ["#3", "England", "9.0%", "55.0% Recent Win Rate, 2019 Champions."],
        ["#4", "South Africa", "8.4%", "95.0 African Host Rating (Co-host Advantage), 64.0% Recent Win Rate."],
        ["#5", "Sri Lanka", "7.2%", "57.1% Knockout Win Rate, 1 WC Title."],
        ["#6", "Pakistan", "5.6%", "52.0% Recent Win Rate, 1992 Champions."],
        ["#7", "West Indies", "5.4%", "2 Historical WC Titles."],
        ["#8", "New Zealand", "4.1%", "2 World Cup Finalist Appearances."],
        ["#9", "Afghanistan", "2.6%", "Solid bowling rating (78.0)."],
        ["#10", "Bangladesh", "2.5%", "38.0% Recent Win Rate."]
    ]

    t_prob = doc.add_table(rows=1, cols=4)
    style_table(t_prob, prob_widths, prob_headers, prob_data)

    p_tbd = doc.add_paragraph()
    p_tbd.paragraph_format.space_before = Pt(8)
    p_tbd.add_run("2027 Actual Winner Designation: ").bold = True
    p_tbd.add_run("Because the 2027 tournament has not taken place, the Actual Winner is designated as 'TBD', and prediction results for both Raj and AI remain 'Pending'.")

    # -------------------------------------------------------------
    # 16. MODEL LIMITATIONS
    # -------------------------------------------------------------
    add_heading_1("16. Model Limitations")

    add_bullet("1. Non-Deterministic Tournament Dynamics: ", "Knockout matches in sports are inherently non-deterministic. A single rain delay, toss decision, or individual performance can alter outcomes regardless of high pre-match probabilities.", "")
    add_bullet("2. Future Squad & Injury Uncertainty: ", "Player injuries, squad restructuring, or form slumps occurring between 2024 and 2027 cannot be anticipated years in advance.", "")
    add_bullet("3. Small Historical Sample Size: ", "Only 11 World Cup editions have taken place between 1983 and 2023, presenting a constrained sample size for statistical modeling.", "")
    add_bullet("4. Synthetic Training Jitter: ", "To train classifiers on 11 historical tournaments, synthetic feature perturbation was applied, which introduces mild modeling assumptions.", "")

    # -------------------------------------------------------------
    # 17. TESTING & VALIDATION
    # -------------------------------------------------------------
    add_heading_1("17. Testing & Verification")

    add_bullet("• 5-Fold Cross-Validation: ", "Executed during ml/train_model.py. Verified accuracy (74.38%), precision (66.46%), recall (46.27%), and F1-score (0.5070).")
    add_bullet("• Historical Backtesting: ", "Evaluated ML predictions across all 11 past tournaments (1983–2023), proving an empirical 72.7% accuracy (8/11 correct).")
    add_bullet("• Backend API Testing: ", "Validated endpoints using curl and python verification scripts (GET /api/prediction/2027, GET /api/history, GET /api/teams).")
    add_bullet("• Frontend Production Build: ", "Ran npm run build in frontend/, confirming a clean compilation in 1.17s with 0 errors.")
    add_bullet("• LocalStorage Persistence Testing: ", "Verified that Raj's manual predictions persist across browser refreshes.")

    # -------------------------------------------------------------
    # 18. INSTALLATION & SETUP
    # -------------------------------------------------------------
    add_heading_1("18. Installation & Setup Guide")

    p = doc.add_paragraph()
    p.add_run("To run the complete project locally on your machine, follow these commands:")

    add_heading_2("1. Backend API & ML Engine Setup")
    p_bcmd = doc.add_paragraph()
    r_bcmd = p_bcmd.add_run(
        "# Navigate to project root\n"
        "cd /Users/rajmalviya/Desktop/cricket-analytics-2027\n\n"
        "# Install Python dependencies\n"
        "pip install -r backend/requirements.txt\n\n"
        "# Train ML Model & generate prediction artifacts\n"
        "python3 ml/train_model.py\n\n"
        "# Start FastAPI Backend Server on port 8000\n"
        "python3 -m uvicorn backend.main:app --host 0.0.0.0 --port 8000 --reload"
    )
    r_bcmd.font.name = 'Consolas'
    r_bcmd.font.size = Pt(9)

    add_heading_2("2. React Frontend Application Setup")
    p_fcmd = doc.add_paragraph()
    r_fcmd = p_fcmd.add_run(
        "# In a new terminal window\n"
        "cd /Users/rajmalviya/Desktop/cricket-analytics-2027/frontend\n\n"
        "# Install Node dependencies\n"
        "npm install\n\n"
        "# Start Vite Dev Server on port 3000\n"
        "npm run dev"
    )
    r_fcmd.font.name = 'Consolas'
    r_fcmd.font.size = Pt(9)

    p_open = doc.add_paragraph()
    p_open.paragraph_format.space_before = Pt(6)
    p_open.add_run("Open ").font.size = Pt(10)
    r_url = p_open.add_run("http://localhost:3000")
    r_url.bold = True
    r_url.font.color.rgb = PRIMARY_COLOR
    p_open.add_run(" in your browser to interact with the application!").font.size = Pt(10)

    # -------------------------------------------------------------
    # 19. DEPLOYMENT
    # -------------------------------------------------------------
    add_heading_1("19. Deployment Guide")

    add_heading_2("Frontend Deployment (Vercel)")
    add_bullet("1. Push project to GitHub: ", "git remote add origin https://github.com/malviyaraj985-glitch/cricket-analytics-2027.git && git push -u origin main", "")
    add_bullet("2. Connect Vercel: ", "Import repo on vercel.com, set Root Directory to frontend, and click Deploy.", "")

    add_heading_2("Backend Deployment (Render)")
    add_bullet("1. Connect Render: ", "Create Web Service on render.com, point to backend directory.", "")
    add_bullet("2. Configure Command: ", "Set start command to python3 -m uvicorn backend.main:app --host 0.0.0.0 --port $PORT.", "")

    # -------------------------------------------------------------
    # 20. FUTURE SCOPE
    # -------------------------------------------------------------
    add_heading_1("20. Future Scope")

    add_bullet("1. Live Ball-by-Ball Data Integration: ", "Connecting ICC live webhooks to update team probabilities dynamically during ongoing series.")
    add_bullet("2. Individual Player Elo Ratings: ", "Incorporating player-level batting/bowling Elo ratings and injury status tracking.")
    add_bullet("3. Monte Carlo Tournament Simulator: ", "Running 10,000 automated tournament match simulations to compute knockout bracket paths.")
    add_bullet("4. Real-time Weather & Pitch Models: ", "Integrating atmospheric and pitch condition APIs for specific South African venues.")

    # -------------------------------------------------------------
    # 21. CONCLUSION
    # -------------------------------------------------------------
    add_heading_1("21. Conclusion")

    p = doc.add_paragraph()
    p.add_run("The ")
    p.add_run("ODI World Cup 2027 AI Analytics & Prediction System").bold = True
    p.add_run(" successfully combines 40 years of empirical ICC World Cup data, feature engineering, machine learning modeling, REST API engineering, and responsive React frontend development. By pairing machine learning predictions with the interactive ")
    p.add_run("Raj vs AI").bold = True
    p.add_run(" feature, the system achieves a 72.7% historical backtesting accuracy score while delivering an engaging, quantitative sports analytics platform suitable for academic presentation and real-world exploration.")

    # -------------------------------------------------------------
    # 22. REFERENCES / DATA SOURCES
    # -------------------------------------------------------------
    add_heading_1("22. References & Data Sources")

    add_bullet("1. Official ICC World Cup Archives: ", "https://www.icc-cricket.com/tournaments/mens-cricket-world-cup")
    add_bullet("2. ESPNcricinfo Statsguru Database: ", "https://stats.espncricinfo.com/ci/engine/stats/index.html")
    add_bullet("3. ICC Men's ODI Team Rankings: ", "https://www.icc-cricket.com/rankings/mens/team-rankings/odi")
    add_bullet("4. Scikit-Learn ML Documentation: ", "https://scikit-learn.org/stable/")
    add_bullet("5. FastAPI Backend Documentation: ", "https://fastapi.tiangolo.com/")

    # -------------------------------------------------------------
    # 23. AUTHOR BLOCK
    # -------------------------------------------------------------
    add_heading_1("23. Author & Project Metadata")

    p_auth = doc.add_paragraph()
    p_auth.paragraph_format.space_before = Pt(6)
    
    r_a1 = p_auth.add_run("Developed by: ")
    r_a1.bold = True
    p_auth.add_run("Raj Malviya\n")
    
    r_a2 = p_auth.add_run("Project Title: ")
    r_a2.bold = True
    p_auth.add_run("ODI World Cup 2027 AI Analytics & Prediction System — Raj vs AI\n")

    r_a3 = p_auth.add_run("Repository Path: ")
    r_a3.bold = True
    p_auth.add_run("/Users/rajmalviya/Desktop/cricket-analytics-2027\n")

    r_a4 = p_auth.add_run("GitHub Remote: ")
    r_a4.bold = True
    p_auth.add_run("https://github.com/malviyaraj985-glitch/cricket-analytics-2027.git\n")

    # Save Document
    target_path = "/Users/rajmalviya/Desktop/cricket-analytics-2027/ODI_World_Cup_2027_Project_Documentation.docx"
    doc.save(target_path)
    print(f"Successfully generated documentation at: {target_path}")

if __name__ == "__main__":
    create_document()
